"""
Legal Brain — Full AI Legal Assistant Router
Capabilities:
1. General legal Q&A (public, no auth required) — GPT-5.4
2. Case-context chat (authenticated) — GPT-5.4 for paid, GPT-5.4 for free
3. Task management — create/schedule tasks, set reminders, auto-alerts
4. Email & communication — draft & send emails via VPS SMTP
5. Document drafting — generate motions, demand letters, discovery requests
6. Research & analysis — case law citations, jurisdiction-specific answers
7. Calendar & reminders — daily briefing, deadline alerts, automated emails
8. Smart proactive suggestions — missing docs, approaching deadlines
"""
from fastapi import APIRouter, HTTPException, Depends
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import BaseModel
from typing import Optional, List
import uuid
import os
import json
import logging
import asyncio
from datetime import datetime, timedelta, timezone

from app.database import get_db
from app.utils.auth import get_current_user, generate_id, decode_token
from app.utils.model_router import get_model_for_task, get_model_for_user_task
from app.utils.credits import credit_gate, deduct_credits

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/legal-brain", tags=["legal-brain"])

security_optional = HTTPBearer(auto_error=False)

# ═══════════════════════════════════════════════════════════
# REQUEST MODELS
# ═══════════════════════════════════════════════════════════

class PublicQuestion(BaseModel):
    question: str
    conversation_id: Optional[str] = None  # Continue existing conversation

class AuthenticatedMessage(BaseModel):
    content: str
    case_id: Optional[str] = None

class TaskCreateRequest(BaseModel):
    case_id: str
    title: str
    description: Optional[str] = None
    due_date: Optional[str] = None
    priority: Optional[str] = "medium"

class EmailDraftRequest(BaseModel):
    case_id: Optional[str] = None
    to_email: str
    subject: str
    context: str  # What the user wants in the email

class DocumentDraftRequest(BaseModel):
    case_id: Optional[str] = None
    document_type: str  # motion_to_compel, demand_letter, discovery_request, etc.
    context: str  # Key facts and requirements

class ReminderRequest(BaseModel):
    case_id: Optional[str] = None
    title: str
    remind_at: str  # ISO datetime
    remind_via: Optional[str] = "email"  # email, notification, both
    notes: Optional[str] = None


# ═══════════════════════════════════════════════════════════
# DB INIT — called from init_db
# ═══════════════════════════════════════════════════════════

LEGAL_BRAIN_TABLES_SQL = """
CREATE TABLE IF NOT EXISTS legal_brain_conversations (
    id TEXT PRIMARY KEY,
    user_id TEXT,
    case_id TEXT,
    is_public INTEGER DEFAULT 0,
    created_at TEXT NOT NULL,
    updated_at TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS legal_brain_messages (
    id TEXT PRIMARY KEY,
    conversation_id TEXT NOT NULL,
    role TEXT NOT NULL CHECK(role IN ('user', 'assistant', 'system')),
    content TEXT NOT NULL,
    metadata_json TEXT DEFAULT '{}',
    created_at TEXT NOT NULL,
    FOREIGN KEY (conversation_id) REFERENCES legal_brain_conversations(id)
);

CREATE TABLE IF NOT EXISTS reminders (
    id TEXT PRIMARY KEY,
    user_id TEXT NOT NULL,
    tenant_id TEXT NOT NULL,
    case_id TEXT,
    title TEXT NOT NULL,
    notes TEXT DEFAULT '',
    remind_at TEXT NOT NULL,
    remind_via TEXT DEFAULT 'email',
    status TEXT DEFAULT 'pending' CHECK(status IN ('pending', 'sent', 'dismissed')),
    created_at TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS scheduled_emails (
    id TEXT PRIMARY KEY,
    user_id TEXT NOT NULL,
    tenant_id TEXT NOT NULL,
    case_id TEXT,
    to_email TEXT NOT NULL,
    subject TEXT NOT NULL,
    body_html TEXT NOT NULL,
    status TEXT DEFAULT 'draft' CHECK(status IN ('draft', 'approved', 'sent', 'failed')),
    send_at TEXT,
    sent_at TEXT,
    created_at TEXT NOT NULL
);

CREATE TABLE IF NOT EXISTS ai_generated_drafts (
    id TEXT PRIMARY KEY,
    user_id TEXT NOT NULL,
    tenant_id TEXT NOT NULL,
    case_id TEXT,
    document_type TEXT NOT NULL,
    title TEXT NOT NULL,
    content TEXT NOT NULL,
    status TEXT DEFAULT 'generated' CHECK(status IN ('generated', 'editing', 'finalized')),
    created_at TEXT NOT NULL
);
"""


def init_legal_brain_tables():
    """Initialize Legal Brain tables. Called from database.init_db()."""
    with get_db() as db:
        db.executescript(LEGAL_BRAIN_TABLES_SQL)
        # Add title column if it doesn't exist (migration for existing databases)
        try:
            db.execute("ALTER TABLE legal_brain_conversations ADD COLUMN title TEXT DEFAULT ''")
        except Exception:
            pass  # Column already exists


def _make_chat_title(message: str) -> str:
    """Clean, document-style fallback title from the first user message."""
    import re as _re
    text = message.strip().replace('\n', ' ')
    text = _re.sub(
        r'^(please\s+|can you\s+|could you\s+|help me\s+|i need to\s+|i need\s+|'
        r'i want to\s+|i want\s+|how do i\s+|what is\s+|what are\s+|'
        r'explain\s+|tell me\s+|analyze\s+|review\s+|draft\s+|write\s+|create\s+)',
        '', text, flags=_re.IGNORECASE
    )
    text = _re.sub(r'\s+', ' ', text).strip()
    words = text.split()
    title = ' '.join(words[:7])
    if len(words) > 7:
        title += '…'
    title = title[0].upper() + title[1:] if title else 'Legal Chat'
    return title[:80]


def _generate_ai_title(conv_id: str, first_message: str) -> None:
    """Background task: ask GPT for a concise 4-6 word conversation title, store in DB."""
    try:
        client = _get_openai_client()
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content":
                    "You name legal chat conversations. Reply with ONLY a concise 4-6 word title "
                    "(no quotes, no punctuation at the end). Examples: "
                    "'Motion to Dismiss Analysis', 'Contract Clause Review', "
                    "'Employment Dispute Strategy', 'Custody Agreement Draft'."},
                {"role": "user", "content": f"Name this conversation based on the first message:\n\n{first_message[:400]}"}
            ],
            max_completion_tokens=20,
            temperature=0.3,
        )
        title = (resp.choices[0].message.content or "").strip().strip('"').strip("'")
        if title:
            from app.database import get_db
            with get_db() as db:
                db.execute(
                    "UPDATE legal_brain_conversations SET title=? WHERE id=?",
                    (title[:80], conv_id)
                )
    except Exception as e:
        logger.warning("[TITLE] AI title generation failed for %s: %s", conv_id, e)


# ═══════════════════════════════════════════════════════════
# HELPER: OpenAI call
# ═══════════════════════════════════════════════════════════

def _get_openai_client():
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(503, "AI service temporarily unavailable")
    from openai import OpenAI
    return OpenAI(api_key=api_key)


def _call_openai(system_prompt: str, user_message: str, model: str = "gpt-5.4",
                 history: list = None, temperature: float = 1.0, max_tokens: int = 16000) -> str:
    """Call OpenAI with system prompt and user message. Returns text response."""
    client = _get_openai_client()
    messages = [{"role": "system", "content": system_prompt}]
    if history:
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": user_message})

    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_completion_tokens=max_tokens,
        )
        return (response.choices[0].message.content or "").strip()
    except Exception as e:
        logger.error(f"[LEGAL BRAIN] OpenAI error: {e}")
        raise HTTPException(503, f"AI service error: {str(e)}")



def _serpapi_search(query: str, num: int = 6) -> str:
    """Fetch live web results from SerpAPI and return formatted context string."""
    import requests as _req
    key = os.environ.get("SERPAPI_KEY", "")
    if not key:
        return ""
    try:
        r = _req.get(
            "https://serpapi.com/search",
            params={"q": query, "api_key": key, "num": num, "engine": "google"},
            timeout=8,
        )
        data = r.json()
        results = data.get("organic_results", [])
        if not results:
            return ""
        lines = [f"[LIVE WEB SEARCH RESULTS for: \"{query}\"]\n"]
        for i, res in enumerate(results[:num], 1):
            title   = res.get("title",   "")
            snippet = res.get("snippet", "")
            link    = res.get("link",    "")
            source  = res.get("source",  "")
            lines.append(f"{i}. {title}" + (f" ({source})" if source else ""))
            if snippet:
                lines.append(f"   {snippet}")
            if link:
                lines.append(f"   {link}")
            lines.append("")
        return "\n".join(lines)
    except Exception as e:
        logger.warning(f"[SEARCH] SerpAPI error: {e}")
        return ""


def _call_openai_with_search(system_prompt: str, user_message: str, model: str = "gpt-5.4",
                              history: list = None, temperature: float = 1.0, max_tokens: int = 16000) -> str:
    """Fetch live web results via SerpAPI then call OpenAI with enriched context."""
    search_context = _serpapi_search(user_message)
    enriched_prompt = system_prompt
    if search_context:
        enriched_prompt = (
            system_prompt
            + "\n\n---\n"
            + search_context
            + "\nUSE the search results above to answer the user's question with current, accurate information. "
            + "Cite the sources and URLs where relevant.\n---"
        )
    client = _get_openai_client()
    messages = [{"role": "system", "content": enriched_prompt}]
    if history:
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": user_message})
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
        max_completion_tokens=max_tokens,
    )
    return (response.choices[0].message.content or "").strip()


def _call_openai_json(system_prompt: str, user_message: str, model: str = "gpt-5.4",
                      temperature: float = 1.0, max_tokens: int = 16000) -> dict:
    """Call OpenAI and parse JSON response."""
    client = _get_openai_client()
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message},
    ]
    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            max_completion_tokens=max_tokens,
            response_format={"type": "json_object"},
        )
        text = (response.choices[0].message.content or "").strip()
        # Strip markdown fences if present
        if text.startswith("```"):
            text = text.split("\n", 1)[1] if "\n" in text else text[3:]
            if text.endswith("```"):
                text = text[:-3]
        return json.loads(text)
    except json.JSONDecodeError as e:
        logger.error(f"[LEGAL BRAIN] JSON parse error: {e}")
        return {"error": "Failed to parse AI response"}
    except Exception as e:
        logger.error(f"[LEGAL BRAIN] OpenAI error: {e}")
        raise HTTPException(503, f"AI service error: {str(e)}")


def _get_legal_db_context(question: str, jurisdiction: str = None):
    """
    RAG-powered semantic search over verified jurisdiction documents.
    Returns (context_string, citations_list).
    Falls back to keyword search if embeddings not yet built.
    """
    try:
        from app.utils.rag_engine import build_rag_context, _ensure_embedding_column
        with get_db() as db:
            _ensure_embedding_column(db)
            context, citations = build_rag_context(db, question, jurisdiction=jurisdiction)
            for c in citations:
                try:
                    db.execute(
                        "UPDATE jurisdiction_documents SET usage_count = usage_count + 1 WHERE title = ?",
                        (c["title"],)
                    )
                except Exception:
                    pass
            return context, citations
    except Exception as e:
        logger.warning(f"[LEGAL BRAIN RAG] Context lookup failed: {e}")
        return "", []

def _search_courtlistener(question: str, max_results: int = 15) -> str:
    """Search CourtListener's free API for real, verified case law.
    Returns formatted case summaries with real citations, courts, and dates."""
    import requests
    import re

    # Extract meaningful search terms from the question
    # Remove common filler words to build a better search query
    stop_words = {
        'i', 'me', 'my', 'we', 'our', 'you', 'your', 'he', 'she', 'it', 'they',
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
        'of', 'with', 'by', 'from', 'is', 'are', 'was', 'were', 'be', 'been',
        'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
        'could', 'should', 'may', 'might', 'shall', 'can', 'need', 'all',
        'give', 'me', 'find', 'get', 'show', 'list', 'what', 'which', 'who',
        'how', 'when', 'where', 'why', 'that', 'this', 'these', 'those',
        'related', 'regarding', 'about', 'case', 'cases', 'law', 'laws',
        'legal', 'court', 'please', 'want', 'like', 'also', 'any', 'some',
    }
    words = re.findall(r'[a-zA-Z]+', question.lower())
    search_terms = [w for w in words if w not in stop_words and len(w) > 2]

    if not search_terms:
        return ""

    # Detect jurisdiction from question for court filtering
    court_filter = ""
    q_lower = question.lower()
    # Map common jurisdiction mentions to CourtListener court IDs
    jurisdiction_map = {
        'new jersey': 'nj,njsuperctappdiv,njsuperctlawdiv,njsupercteqdiv,njch,njsuperctchandiv',
        'new york': 'ny,nyappdiv,nyappterm,nysupct,nycivct,nyfamct',
        'california': 'cal,calctapp,calag',
        'texas': 'tex,texapp,texcrimapp',
        'florida': 'fla,fladistctapp',
        'illinois': 'ill,illappct',
        'pennsylvania': 'pa,pasuperct,pacommwct',
        'georgia': 'ga,gactapp',
        'ohio': 'ohio,ohioctapp,ohioctcl',
        'michigan': 'mich,michctapp',
        'virginia': 'va,vactapp',
        'massachusetts': 'mass,massappct',
        'maryland': 'md,mdctspecapp',
        'north carolina': 'nc,ncctapp',
        'washington': 'wash,washctapp',
        'connecticut': 'conn,connappct,connsuperct',
        'colorado': 'colo,coloctapp',
        'minnesota': 'minn,minnctapp',
        'wisconsin': 'wis,wisctapp',
        'oregon': 'or,orctapp',
        'louisiana': 'la,lactapp',
        'alabama': 'ala,alacivapp,alacrimapp',
        'south carolina': 'sc,scctapp',
        'kentucky': 'ky,kyctapp',
        'indiana': 'ind,indctapp,indtc',
        'tennessee': 'tenn,tennctapp,tenncrimapp',
        'arizona': 'ariz,arizctapp',
        'iowa': 'iowa,iowactapp',
        'arkansas': 'ark,arkctapp',
        'missouri': 'mo,moctapp',
        'nebraska': 'neb,nebctapp',
        'delaware': 'del,delch,delsuperct,delctcompl',
        'uganda': '',
        'federal': 'scotus,ca1,ca2,ca3,ca4,ca5,ca6,ca7,ca8,ca9,ca10,ca11,cadc,cafc',
    }

    for jurisdiction_name, court_ids in jurisdiction_map.items():
        if jurisdiction_name in q_lower:
            court_filter = court_ids
            break

    # Build search query - use spaces for natural query syntax
    search_query = " ".join(search_terms[:8])  # Limit to 8 key terms

    try:
        params = {
            "q": search_query,
            "type": "o",  # opinions
            "page_size": max_results,
            "order_by": "score desc",
        }
        if court_filter:
            params["court"] = court_filter

        response = requests.get(
            "https://www.courtlistener.com/api/rest/v4/search/",
            params=params,
            timeout=15,
            headers={"User-Agent": "LitigationSpace-LegalBrain/1.0"}
        )

        if response.status_code != 200:
            logger.warning(f"[LEGAL BRAIN] CourtListener API returned {response.status_code}")
            return ""

        data = response.json()
        results = data.get("results", [])
        total_count = data.get("count", 0)

        if not results:
            return ""

        case_entries = []
        case_entries.append(f"COURTLISTENER SEARCH RESULTS ({total_count} total matches, showing top {len(results)}):\n")

        for i, r in enumerate(results, 1):
            case_name = r.get("caseName", "Unknown Case")
            citations = r.get("citation", [])
            citation_str = "; ".join(citations) if citations else "No official citation"
            court_name = r.get("court", "Unknown Court")
            date_filed = r.get("dateFiled", "Unknown date")
            docket_num = r.get("docketNumber", "")
            cite_count = r.get("citeCount", 0)
            url = "https://www.courtlistener.com" + r.get("absolute_url", "")
            snippet = r.get("snippet", "")

            entry = f"Case {i}: {case_name}\n"
            entry += f"  Citation: {citation_str}\n"
            entry += f"  Court: {court_name}\n"
            entry += f"  Date Filed: {date_filed}\n"
            if docket_num:
                entry += f"  Docket: {docket_num}\n"
            entry += f"  Times Cited: {cite_count}\n"
            entry += f"  CourtListener URL: {url}\n"
            if snippet:
                # Clean HTML tags from snippet
                clean_snippet = re.sub(r'<[^>]+>', '', snippet)
                entry += f"  Excerpt: {clean_snippet[:600]}\n"

            case_entries.append(entry)

        return "\n".join(case_entries)

    except Exception as e:
        logger.warning(f"[LEGAL BRAIN] CourtListener search error: {e}")
        return ""


def _get_user_model(user_row) -> str:
    """Determine which model to use based on user's subscription tier.
    Delegates to central model router."""
    if not user_row:
        return get_model_for_task("legal_brain_public")
    status = user_row.get("status", "LOCKED") if isinstance(user_row, dict) else (user_row["status"] if user_row else "LOCKED")
    is_paid = status in ("ACTIVE", "PREMIUM", "PRO", "READY")
    return get_model_for_user_task("legal_brain_chat", "legal_brain_public", is_paid)


def _get_case_context(db, case_id: str, tenant_id: str) -> dict:
    """Build comprehensive case context for AI."""
    case = db.execute("SELECT * FROM cases WHERE id = ? AND tenant_id = ?", (case_id, tenant_id)).fetchone()
    if not case:
        return {}
    case_dict = dict(case)

    tasks = [dict(r) for r in db.execute(
        "SELECT title, status, due_date, priority FROM tasks WHERE case_id = ? AND tenant_id = ? ORDER BY due_date ASC",
        (case_id, tenant_id)
    ).fetchall()]

    docs = [dict(r) for r in db.execute(
        "SELECT filename, category FROM documents WHERE case_id = ? AND tenant_id = ?",
        (case_id, tenant_id)
    ).fetchall()]

    witnesses = [dict(r) for r in db.execute(
        "SELECT name, witness_type, deposition_date FROM witnesses WHERE case_id = ? AND tenant_id = ?",
        (case_id, tenant_id)
    ).fetchall()]

    discovery = [dict(r) for r in db.execute(
        "SELECT item_description, status, date_due FROM discovery_items WHERE case_id = ? AND tenant_id = ?",
        (case_id, tenant_id)
    ).fetchall()]

    drafts = [dict(r) for r in db.execute(
        "SELECT title, document_type, status FROM legal_drafts WHERE case_id = ? AND tenant_id = ?",
        (case_id, tenant_id)
    ).fetchall()]

    return {
        "case": case_dict,
        "tasks": tasks,
        "documents": docs,
        "witnesses": witnesses,
        "discovery": discovery,
        "drafts": drafts,
    }


def _format_case_context(ctx: dict) -> str:
    """Format case context into a string for the system prompt."""
    if not ctx or not ctx.get("case"):
        return "No case context available."

    c = ctx["case"]
    lines = [
        f'CASE: "{c.get("title", "Untitled")}"',
        f'Type: {(c.get("case_type", "") or "").replace("_", " ").title()}',
    ]
    if c.get("court"):
        lines.append(f'Court: {c["court"]}')
    if c.get("judge"):
        lines.append(f'Judge: {c["judge"]}')
    if c.get("filing_deadline"):
        lines.append(f'Filing Deadline: {c["filing_deadline"]}')
    if c.get("trial_date"):
        lines.append(f'Trial Date: {c["trial_date"]}')
    if c.get("opposing_party"):
        lines.append(f'Opposing Party: {c["opposing_party"]}')
    if c.get("client_name"):
        lines.append(f'Client: {c["client_name"]}')

    tasks = ctx.get("tasks", [])
    pending = [t for t in tasks if t.get("status") == "pending"]
    overdue = [t for t in tasks if t.get("due_date") and t.get("status") != "completed"
               and t["due_date"] < datetime.now(timezone.utc).isoformat()]
    completed = [t for t in tasks if t.get("status") == "completed"]
    total = len(tasks)
    pct = round(len(completed) / total * 100) if total > 0 else 0
    lines.append(f'\nProgress: {pct}% ({len(completed)}/{total} tasks)')

    if overdue:
        lines.append(f'\n⚠️ OVERDUE TASKS ({len(overdue)}):')
        for t in overdue[:5]:
            lines.append(f'  - {t["title"]} (due: {t.get("due_date", "?")})')

    if pending:
        lines.append(f'\nPending Tasks ({len(pending)}):')
        for t in pending[:10]:
            lines.append(f'  - {t["title"]} (due: {t.get("due_date", "N/A")}, priority: {t.get("priority", "medium")})')

    docs = ctx.get("documents", [])
    if docs:
        lines.append(f'\nDocuments ({len(docs)}):')
        for d in docs[:10]:
            lines.append(f'  - {d["filename"]} ({d.get("category", "general")})')
    else:
        lines.append('\nDocuments: None uploaded yet')

    witnesses = ctx.get("witnesses", [])
    if witnesses:
        lines.append(f'\nWitnesses ({len(witnesses)}):')
        for w in witnesses[:10]:
            lines.append(f'  - {w["name"]} ({w.get("witness_type", "fact")})')

    discovery = ctx.get("discovery", [])
    if discovery:
        pending_d = [d for d in discovery if d.get("status") == "pending"]
        overdue_d = [d for d in discovery if d.get("status") == "overdue"]
        lines.append(f'\nDiscovery: {len(discovery)} items, {len(pending_d)} pending, {len(overdue_d)} overdue')

    drafts = ctx.get("drafts", [])
    if drafts:
        lines.append(f'\nDrafts ({len(drafts)}):')
        for d in drafts[:5]:
            lines.append(f'  - {d["title"]} ({d.get("document_type", "motion")}) — {d.get("status", "draft")}')

    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════
# 1. PUBLIC LEGAL Q&A (no auth required)
# ═══════════════════════════════════════════════════════════

@router.post("/ask")
async def public_ask(req: PublicQuestion, credentials: HTTPAuthorizationCredentials = Depends(security_optional)):
    """Public legal Q&A — free-tier response via model router."""
    model = get_model_for_task("legal_brain_public")
    user_id = None

    # If authenticated, track user
    if credentials and credentials.credentials:
        try:
            user_data = decode_token(credentials.credentials)
            user_id = user_data.get("sub")
        except Exception:
            pass

    # Fetch relevant legal DB context via RAG vector search
    legal_db_context, rag_citations = _get_legal_db_context(req.question)

    # Search CourtListener for real, verified case law
    courtlistener_results = _search_courtlistener(req.question, max_results=20)
    logger.info(f"[LEGAL BRAIN] CourtListener returned {len(courtlistener_results)} chars for query: {req.question[:80]}")

    current_date = datetime.now().strftime("%B %d, %Y")

    system_prompt = f"""You are the **LitigationSpace Legal Brain** — the most advanced AI legal research and analysis engine available. You are powered by **GPT-5.4** and backed by **CourtListener's verified case law database** containing millions of real court opinions.

Today's date is {current_date}. You have REAL-TIME WEB SEARCH capability — use it actively.

When users ask about current news, recent elections, new legislation, recent court decisions, current statistics, or any time-sensitive information — SEARCH THE WEB and provide current, accurate facts. Never say "I cannot access real-time information" or "my knowledge cuts off at" — you have live web search. Use it and report what you find.

You deliver **partner-level legal research memos** that litigation professionals can rely on. You are NOT a lawyer and do not provide legal advice, but your analysis must be as thorough, detailed, and well-cited as a top-tier law firm's research product.

---

## ABSOLUTE RULE #1 — CITATION INTEGRITY (NON-NEGOTIABLE)

You have been provided **REAL, VERIFIED case law** from CourtListener's database below. These are actual court opinions from a verified legal database.

**YOU MUST:**
- **USE the CourtListener cases below as your PRIMARY citation source** — these are verified real cases
- **Copy the EXACT case name, citation, court, and date** from the CourtListener results — do NOT modify them
- **Include the CourtListener URL** for each case so users can read the full opinion
- You may ALSO cite well-known U.S. Supreme Court landmarks and other widely-cited decisions you are highly confident about
- For legal principles where you know the rule but lack a verified citation: state the principle clearly and note "(established by case law; verify specific citation via Westlaw/LexisNexis)"

**YOU MUST NEVER:**
- Fabricate case names, volume numbers, reporter pages, or citations
- Invent cases that sound plausible but don't exist
- Modify or "improve" citations from the CourtListener results
- Say "I cannot provide specific case citations" when CourtListener results are provided below — USE THEM

{f"""### Verified Legal Database Sources (RAG)
The following are semantically matched excerpts from LitigationSpace's verified jurisdiction database:

{legal_db_context}

**Sources above have been retrieved via vector similarity search and are verified legal documents.**""" if legal_db_context else ""}

{f"### VERIFIED CASE LAW FROM COURTLISTENER DATABASE{chr(10)}*(These are REAL cases retrieved from CourtListener's verified database of millions of court opinions. You MUST cite these cases in your response.)*{chr(10)}{chr(10)}{courtlistener_results}" if courtlistener_results else "*No CourtListener results available for this specific query. Only cite cases you are highly confident are real. For any case you cite, include the full citation and note whether it is a landmark case or requires verification.*"}

---

## RESPONSE FORMAT & STYLE

Use **rich markdown formatting** for professional, readable presentation:
- Use `##` and `###` headings to organize major sections
- Use **bold** for case names, key terms, and important concepts
- Use *italics* for emphasis and legal Latin terms
- Use numbered lists (1. 2. 3.) for sequential steps or ranked items
- Use bullet points for non-sequential lists
- Use blockquotes for key holdings or important quotes from cases
- Use --- horizontal rules to separate major sections
- Include clickable links to CourtListener cases
- Keep paragraphs short -- maximum 3 sentences each
- Never use walls of text -- break everything into scannable sections
- Be precise and concise -- lawyers value brevity
- End every response with a **Bottom Line / Key Takeaway** section summarising the critical points

## DEPTH & THOROUGHNESS REQUIREMENTS

**Your responses must be EXHAUSTIVE and COMPREHENSIVE:**
- **Minimum 3,000-5,000 words** for substantive legal questions about case law
- **Minimum 2,000-3,000 words** for general legal questions
- **No surface-level summaries** — every point must be analyzed in depth
- For each case cited: provide the **full citation, court, year, key facts, holding, reasoning, and strategic significance**
- Discuss **practical implications** — how does this law play out in real disputes?
- Address **counterarguments** and how to respond to them
- Note **jurisdictional variations** and differences across courts
- Include **strategic litigation tips** for practitioners

## RESPONSE STRUCTURE

For case law research questions, structure your response as:

1. **Executive Summary** — Direct answer to the question (2-3 paragraphs)
2. **Governing Legal Framework** — Applicable statutes, rules, and standards with section numbers
3. **Key Case Law Analysis** — Organized by legal themes/subtopics, NOT a flat list
   - For each case: full citation, court, facts, holding, reasoning, and significance
   - Distinguish between binding precedent and persuasive authority
4. **Practical Implications & Strategy** — How practitioners can use this law
5. **Potential Counterarguments & Responses** — Anticipate opposing arguments
6. **Jurisdictional Considerations** — Circuit splits, state variations, evolving law
7. **Key Takeaways** — Bullet-point summary of the most critical points
8. **Recommended Next Steps** — Concrete, actionable guidance
9. **Disclaimer** — Note that citations should be verified via Westlaw/LexisNexis

For document review questions:
1. Identify the document type and purpose
2. Analyze **EVERY section, clause, and paragraph** — skip nothing
3. Flag problematic clauses, ambiguities, and missing provisions
4. Highlight favorable and unfavorable terms
5. Provide risk assessment and recommended modifications

## CRITICAL — DOCUMENT REDRAFTING & REWRITING REQUESTS

When a user asks you to **redraft**, **rewrite**, **revise**, **improve**, **edit**, or **reformat** a document (motion, brief, letter, contract, etc.):

**YOU MUST produce the FULL, COMPLETE redrafted document — not guidance on how to redraft it.**

- Output the **entire rewritten document** from start to finish — every section, every paragraph, every sentence
- Do NOT say "here's how you should revise it" or "consider making these changes" — actually MAKE the changes and produce the final product
- Maintain proper legal document formatting with headings, numbered paragraphs, and citations
- Use **rich markdown formatting** for the redrafted document
- If the user provides a document and asks you to strengthen it, fix it, or improve it — produce the COMPLETE improved version
- Minimum length: match or exceed the original document length
- The user should be able to copy your output and use it directly

**If the user needs to generate a brand-new document from scratch** (e.g., "draft me a motion to compel"), mention that LitigationSpace's **Drafting Room** feature is specifically designed for generating court-ready documents with proper formatting, and they can find it in the sidebar navigation. But if they insist or ask you directly, produce the full document anyway.

---

Write like a **senior litigation partner** preparing a comprehensive research memo for a major case. Your analysis should be so thorough that a lawyer reading it feels they have a complete understanding of the legal landscape. **Never be brief. Never be superficial. Exhaust the topic.**"""

    # Get or create conversation
    now = datetime.now(timezone.utc).isoformat()
    history = []
    conv_id = req.conversation_id

    if conv_id:
        # Load existing conversation history for context
        with get_db() as db:
            existing = db.execute(
                "SELECT role, content FROM legal_brain_messages WHERE conversation_id = ? ORDER BY created_at ASC",
                (conv_id,)
            ).fetchall()
            history = [{"role": dict(m)["role"], "content": dict(m)["content"]} for m in existing]
    else:
        conv_id = generate_id()

    # Attempt web search tool; fall back silently if model does not support it
    try:
        ai_response = await asyncio.get_event_loop().run_in_executor(
            None, lambda: _call_openai_with_search(system_prompt, req.question, model=model, history=history, max_tokens=16000)
        )
    except Exception:
        ai_response = await asyncio.get_event_loop().run_in_executor(
            None, lambda: _call_openai(system_prompt, req.question, model=model, history=history, max_tokens=16000)
        )

    # Save conversation
    is_new = not req.conversation_id
    with get_db() as db:
        if is_new:
            db.execute(
                "INSERT INTO legal_brain_conversations (id, user_id, case_id, is_public, title, created_at, updated_at) VALUES (?, ?, NULL, 1, ?, ?, ?)",
                (conv_id, user_id, _make_chat_title(req.question), now, now)
            )
            import threading as _thr
            _thr.Thread(target=_generate_ai_title, args=(conv_id, req.question), daemon=True).start()
        else:
            db.execute("UPDATE legal_brain_conversations SET updated_at = ? WHERE id = ?", (now, conv_id))
        db.execute(
            "INSERT INTO legal_brain_messages (id, conversation_id, role, content, created_at) VALUES (?, ?, 'user', ?, ?)",
            (generate_id(), conv_id, req.question, now)
        )
        db.execute(
            "INSERT INTO legal_brain_messages (id, conversation_id, role, content, created_at) VALUES (?, ?, 'assistant', ?, ?)",
            (generate_id(), conv_id, ai_response, now)
        )

    return {
        "conversation_id": conv_id,
        "response": ai_response,
        "model": model,
    }


# Alias used by the public /legal-brain SPA guest interface
@router.post("/public/chat")
async def public_chat_alias(req: PublicQuestion, credentials: HTTPAuthorizationCredentials = Depends(security_optional)):
    """Alias of /ask — called by the guest-facing Legal Brain SPA."""
    return await public_ask(req, credentials)


# ═══════════════════════════════════════════════════════════
# 2. CASE-CONTEXT CHAT (authenticated)
# ═══════════════════════════════════════════════════════════

@router.post("/chat")
async def case_chat(req: AuthenticatedMessage, user=Depends(get_current_user)):
    """Case-context AI chat — full assistant with action capabilities."""
    user_id = user["sub"]
    tenant_id = user["tenant_id"]

    with get_db() as db:
        model, _lb_cost = credit_gate(user_id, "legal_brain_chat", db)
        _lb_status = db.execute("SELECT subscription_status FROM users WHERE id=?", (user_id,)).fetchone()["subscription_status"]

        # Build case context if case_id provided
        case_context_str = ""
        if req.case_id:
            ctx = _get_case_context(db, req.case_id, tenant_id)
            case_context_str = _format_case_context(ctx)

        # Get recent conversation history
        recent = db.execute("""
            SELECT lbm.role, lbm.content FROM legal_brain_messages lbm
            JOIN legal_brain_conversations lbc ON lbm.conversation_id = lbc.id
            WHERE lbc.user_id = ? AND (lbc.case_id = ? OR lbc.case_id IS NULL)
            ORDER BY lbm.created_at DESC LIMIT 10
        """, (user_id, req.case_id)).fetchall()
        history = [dict(m) for m in reversed(recent)]

        # Get all cases for this user (for cross-case awareness)
        all_cases = db.execute(
            "SELECT id, title, case_type, status, filing_deadline FROM cases WHERE tenant_id = ? ORDER BY updated_at DESC LIMIT 10",
            (tenant_id,)
        ).fetchall()
        cases_summary = "\n".join([f"- {dict(c)['title']} ({dict(c)['case_type']}) — {dict(c)['status']}" for c in all_cases])

        # Check for approaching deadlines across all cases
        upcoming_deadlines = db.execute("""
            SELECT c.title as case_title, t.title as task_title, t.due_date, t.status
            FROM tasks t JOIN cases c ON t.case_id = c.id
            WHERE t.tenant_id = ? AND t.status IN ('pending', 'in_progress')
            AND t.due_date IS NOT NULL AND t.due_date != ''
            ORDER BY t.due_date ASC LIMIT 15
        """, (tenant_id,)).fetchall()
        deadlines_str = ""
        if upcoming_deadlines:
            deadlines_str = "\n\nUPCOMING DEADLINES:\n"
            for d in upcoming_deadlines:
                dd = dict(d)
                deadlines_str += f"- [{dd.get('case_title', '?')}] {dd['task_title']} — due: {dd.get('due_date', '?')} ({dd['status']})\n"

    now_str = datetime.now(timezone.utc).strftime("%A, %B %d, %Y at %I:%M %p UTC")

    system_prompt = f"""You are the LitigationSpace AI Case Navigator — a full-service legal assistant embedded in a litigation management platform. Today is {now_str}.

You have FULL ACCESS to this user's data and can take ACTION on their behalf.

You have REAL-TIME WEB SEARCH capability. When users ask about current news, recent elections, new legislation, recent court decisions, current statistics, or any time-sensitive information — USE web search to fetch up-to-date information. Never say "I cannot access real-time information" or "my knowledge cuts off at" — search for it instead and provide current facts.

{f"CURRENT CASE CONTEXT:{chr(10)}{case_context_str}" if case_context_str else "No specific case selected."}

ALL USER CASES:
{cases_summary if cases_summary else "No cases yet."}
{deadlines_str}

YOUR CAPABILITIES (action-oriented):

1. **Task Management**: Create tasks, set deadlines, mark tasks complete, identify overdue items
   - When the user asks to create a task, respond with a JSON action block: {{"action": "create_task", "case_id": "...", "title": "...", "due_date": "YYYY-MM-DD", "priority": "high/medium/low"}}

2. **Email & Communication**: Draft professional legal emails, ready to send
   - When drafting emails, respond with: {{"action": "draft_email", "to": "...", "subject": "...", "body": "..."}}

3. **Document Drafting**: Generate motions, demand letters, discovery requests
   - When drafting documents, respond with: {{"action": "draft_document", "type": "...", "title": "...", "content": "..."}}

4. **Research & Analysis**: Case law citations, jurisdiction-specific answers, motion analysis
   - ONLY cite cases you are HIGHLY CONFIDENT are real — never fabricate citations
   - If unsure of exact citation, describe the legal principle without a fake cite
   - Reference relevant statutes and procedural rules
   - Connect to the Win Probability Simulator when analyzing motions

5. **Calendar & Reminders**: Track deadlines, set reminders, provide daily briefings
   - When setting reminders: {{"action": "set_reminder", "title": "...", "remind_at": "YYYY-MM-DDTHH:MM", "notes": "..."}}

6. **Smart Suggestions**: Proactively identify missing documents, approaching deadlines, strategic opportunities

RESPONSE FORMAT:
- Use rich markdown formatting: ## headings, **bold**, *italics*, numbered lists, bullet points, blockquotes
- Use numbered lists (1. 2. 3.) for sequential steps or ranked items; bullet points for non-sequential items
- Keep paragraphs short — maximum 3 sentences each
- Use bold headings for each major section; add --- dividers between sections
- Never use walls of text — break everything into scannable sections
- Be precise and concise — lawyers value brevity
- End every response with a **Bottom Line / Key Takeaway** section
- Give THOROUGH, DETAILED, EXHAUSTIVE answers — never brief or superficial. **Minimum 1,500-3,000 words** for substantive questions
- For action responses, include BOTH a thorough human-readable explanation AND a JSON action block wrapped in ```action``` code fences
- Always be specific to the case facts and jurisdiction
- End with concrete next steps
- Be authoritative but note this is AI-assisted analysis, not legal advice

IMPORTANT: When you detect the user wants to CREATE something (task, email, document, reminder), you MUST include the action JSON block so the frontend can execute it.

## CRITICAL — DOCUMENT REDRAFTING & REWRITING

When a user asks you to **redraft**, **rewrite**, **revise**, **improve**, **edit**, or **reformat** a document:

**YOU MUST produce the FULL, COMPLETE redrafted document — not just guidance on how to redraft it.**

- Output the **entire rewritten document** from start to finish
- Do NOT say "here's how you should revise it" — actually MAKE the changes and produce the final product
- Maintain proper legal document formatting
- The user should be able to copy your output and use it directly
- Match or exceed the original document length

**For brand-new documents from scratch**, mention LitigationSpace's **Drafting Room** (in the sidebar) which is purpose-built for generating court-ready documents. But if they insist, produce the full document anyway."""

    try:
        ai_response = await asyncio.get_event_loop().run_in_executor(
            None, lambda: _call_openai_with_search(system_prompt, req.content, model=model, history=history, max_tokens=16000)
        )
    except Exception:
        ai_response = await asyncio.get_event_loop().run_in_executor(
            None, lambda: _call_openai(system_prompt, req.content, model=model, history=history, max_tokens=16000)
        )

    # Parse any action blocks from the response
    actions = _extract_actions(ai_response)

    # Execute any actions
    executed_actions = []
    for action in actions:
        result = await _execute_action(action, user_id, tenant_id, req.case_id)
        if result:
            executed_actions.append(result)

    # Save messages
    now = datetime.now(timezone.utc).isoformat()
    with get_db() as db:
        # Find or create conversation
        conv = db.execute(
            "SELECT id FROM legal_brain_conversations WHERE user_id = ? AND case_id IS ? ORDER BY updated_at DESC LIMIT 1",
            (user_id, req.case_id)
        ).fetchone()
        if conv:
            conv_id = conv["id"]
            db.execute("UPDATE legal_brain_conversations SET updated_at = ? WHERE id = ?", (now, conv_id))
        else:
            conv_id = generate_id()
            db.execute(
                "INSERT INTO legal_brain_conversations (id, user_id, case_id, is_public, title, created_at, updated_at) VALUES (?, ?, ?, 0, ?, ?, ?)",
                (conv_id, user_id, req.case_id, _make_chat_title(req.content), now, now)
            )
            import threading as _thr
            _thr.Thread(target=_generate_ai_title, args=(conv_id, req.question), daemon=True).start()
        db.execute(
            "INSERT INTO legal_brain_messages (id, conversation_id, role, content, metadata_json, created_at) VALUES (?, ?, 'user', ?, '{}', ?)",
            (generate_id(), conv_id, req.content, now)
        )
        db.execute(
            "INSERT INTO legal_brain_messages (id, conversation_id, role, content, metadata_json, created_at) VALUES (?, ?, 'assistant', ?, ?, ?)",
            (generate_id(), conv_id, ai_response, json.dumps({"actions": executed_actions}), now)
        )
        deduct_credits(user_id, _lb_status, _lb_cost, "legal_brain_chat", db)

    return {
        "conversation_id": conv_id,
        "response": ai_response,
        "actions_executed": executed_actions,
    }


def _extract_actions(response: str) -> list:
    """Extract action JSON blocks from AI response."""
    actions = []
    import re
    # Look for ```action ... ``` blocks
    pattern = r'```action\s*\n(.*?)\n```'
    matches = re.findall(pattern, response, re.DOTALL)
    for match in matches:
        try:
            action = json.loads(match.strip())
            if "action" in action:
                actions.append(action)
        except json.JSONDecodeError:
            continue
    # Also look for inline JSON with "action" key
    json_pattern = r'\{[^{}]*"action"\s*:\s*"[^"]*"[^{}]*\}'
    for match in re.finditer(json_pattern, response):
        try:
            action = json.loads(match.group())
            if action not in actions:
                actions.append(action)
        except json.JSONDecodeError:
            continue
    return actions


async def _execute_action(action: dict, user_id: str, tenant_id: str, default_case_id: str = None) -> dict:
    """Execute an AI-generated action."""
    action_type = action.get("action")
    now = datetime.now(timezone.utc).isoformat()

    try:
        if action_type == "create_task":
            task_id = generate_id()
            case_id = action.get("case_id", default_case_id)
            if not case_id:
                return {"action": "create_task", "status": "failed", "reason": "No case selected"}
            with get_db() as db:
                db.execute(
                    "INSERT INTO tasks (id, case_id, tenant_id, title, description, status, due_date, priority, created_at) VALUES (?, ?, ?, ?, ?, 'pending', ?, ?, ?)",
                    (task_id, case_id, tenant_id, action.get("title", "Untitled Task"),
                     action.get("description", ""), action.get("due_date"), action.get("priority", "medium"), now)
                )
            return {"action": "create_task", "status": "success", "task_id": task_id, "title": action.get("title")}

        elif action_type == "draft_email":
            email_id = generate_id()
            with get_db() as db:
                db.execute(
                    "INSERT INTO scheduled_emails (id, user_id, tenant_id, case_id, to_email, subject, body_html, status, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, 'draft', ?)",
                    (email_id, user_id, tenant_id, action.get("case_id", default_case_id),
                     action.get("to", ""), action.get("subject", ""), action.get("body", ""), now)
                )
            return {"action": "draft_email", "status": "success", "email_id": email_id, "subject": action.get("subject")}

        elif action_type == "draft_document":
            doc_id = generate_id()
            with get_db() as db:
                db.execute(
                    "INSERT INTO ai_generated_drafts (id, user_id, tenant_id, case_id, document_type, title, content, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
                    (doc_id, user_id, tenant_id, action.get("case_id", default_case_id),
                     action.get("type", "motion"), action.get("title", "Untitled"), action.get("content", ""), now)
                )
            return {"action": "draft_document", "status": "success", "document_id": doc_id, "title": action.get("title")}

        elif action_type == "set_reminder":
            reminder_id = generate_id()
            with get_db() as db:
                db.execute(
                    "INSERT INTO reminders (id, user_id, tenant_id, case_id, title, notes, remind_at, remind_via, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, 'email', ?)",
                    (reminder_id, user_id, tenant_id, action.get("case_id", default_case_id),
                     action.get("title", "Reminder"), action.get("notes", ""), action.get("remind_at", ""), now)
                )
            return {"action": "set_reminder", "status": "success", "reminder_id": reminder_id, "title": action.get("title")}

    except Exception as e:
        logger.error(f"[LEGAL BRAIN] Action execution error: {e}")
        return {"action": action_type, "status": "failed", "reason": str(e)}

    return {}


# ═══════════════════════════════════════════════════════════
# 3. TASK MANAGEMENT VIA AI
# ═══════════════════════════════════════════════════════════

@router.post("/tasks/create")
async def ai_create_task(req: TaskCreateRequest, user=Depends(get_current_user)):
    """Create a task via the Legal Brain. Also generates smart sub-tasks."""
    task_id = generate_id()
    now = datetime.now(timezone.utc).isoformat()

    with get_db() as db:
        # Verify case access
        case = db.execute("SELECT * FROM cases WHERE id = ? AND tenant_id = ?",
                          (req.case_id, user["tenant_id"])).fetchone()
        if not case:
            raise HTTPException(404, "Case not found")

        db.execute(
            "INSERT INTO tasks (id, case_id, tenant_id, title, description, status, due_date, priority, created_at) VALUES (?, ?, ?, ?, ?, 'pending', ?, ?, ?)",
            (task_id, req.case_id, user["tenant_id"], req.title, req.description or "", req.due_date, req.priority, now)
        )

    # Generate smart sub-task suggestions
    system_prompt = """You are a legal task management AI. Given a legal task, suggest 3-5 specific sub-tasks that would help complete it.
Return JSON: {"sub_tasks": [{"title": "...", "priority": "high/medium/low"}]}"""

    try:
        result = _call_openai_json(system_prompt, f"Task: {req.title}\nCase type: {dict(case).get('case_type', 'litigation')}\nDescription: {req.description or 'N/A'}")
        sub_tasks = result.get("sub_tasks", [])
    except Exception:
        sub_tasks = []

    return {
        "task_id": task_id,
        "title": req.title,
        "suggested_sub_tasks": sub_tasks,
    }


@router.get("/tasks/overdue")
async def get_overdue_tasks(user=Depends(get_current_user)):
    """Get all overdue tasks with AI-generated priority recommendations."""
    now = datetime.now(timezone.utc).isoformat()
    with get_db() as db:
        overdue = db.execute("""
            SELECT t.*, c.title as case_title FROM tasks t
            JOIN cases c ON t.case_id = c.id
            WHERE t.tenant_id = ? AND t.status IN ('pending', 'in_progress')
            AND t.due_date IS NOT NULL AND t.due_date < ?
            ORDER BY t.due_date ASC
        """, (user["tenant_id"], now)).fetchall()
        tasks = [dict(r) for r in overdue]

    if not tasks:
        return {"overdue_tasks": [], "recommendation": "No overdue tasks. You're on track!"}

    # AI recommendation
    task_list = "\n".join([f"- {t['title']} (case: {t.get('case_title', '?')}, due: {t.get('due_date', '?')})" for t in tasks])
    system_prompt = "You are a legal task prioritization AI. Rank these overdue tasks by urgency and provide a brief action plan. Use markdown."
    recommendation = _call_openai(system_prompt, f"Overdue tasks:\n{task_list}", max_tokens=800)

    return {"overdue_tasks": tasks, "recommendation": recommendation}


@router.get("/tasks/upcoming")
async def get_upcoming_tasks(days: int = 7, user=Depends(get_current_user)):
    """Get upcoming tasks within N days with AI prioritization."""
    now = datetime.now(timezone.utc)
    future = (now + timedelta(days=days)).isoformat()
    now_str = now.isoformat()

    with get_db() as db:
        upcoming = db.execute("""
            SELECT t.*, c.title as case_title FROM tasks t
            JOIN cases c ON t.case_id = c.id
            WHERE t.tenant_id = ? AND t.status IN ('pending', 'in_progress')
            AND t.due_date IS NOT NULL AND t.due_date >= ? AND t.due_date <= ?
            ORDER BY t.due_date ASC
        """, (user["tenant_id"], now_str, future)).fetchall()

    return {"upcoming_tasks": [dict(r) for r in upcoming], "period_days": days}


# ═══════════════════════════════════════════════════════════
# 4. EMAIL & COMMUNICATION
# ═══════════════════════════════════════════════════════════

@router.post("/email/draft")
async def ai_draft_email(req: EmailDraftRequest, user=Depends(get_current_user)):
    """Draft a professional legal email using AI."""
    with get_db() as db:
        user_row = db.execute("SELECT full_name, email FROM users WHERE id = ?", (user["sub"],)).fetchone()
        user_status = db.execute("SELECT status FROM users WHERE id = ?", (user["sub"],)).fetchone()
        model = _get_user_model(dict(user_status)) if user_status else get_model_for_task("legal_brain_draft_email")

        case_context = ""
        if req.case_id:
            ctx = _get_case_context(db, req.case_id, user["tenant_id"])
            case_context = _format_case_context(ctx)

    sender_name = dict(user_row)["full_name"] if user_row else "Attorney"
    sender_email = dict(user_row)["email"] if user_row else ""

    system_prompt = f"""You are a professional legal email drafting assistant. Draft a polished, professional legal email.

Sender: {sender_name} ({sender_email})
{f"Case Context: {case_context}" if case_context else ""}

Rules:
- Professional tone appropriate for legal correspondence
- Clear, concise language
- Proper salutation and closing
- Include relevant case references if applicable
- Return JSON: {{"subject": "...", "body_html": "<p>...</p>"}}"""

    result = _call_openai_json(system_prompt,
                                f"To: {req.to_email}\nSubject: {req.subject}\nContext: {req.context}",
                                model=model)

    email_id = generate_id()
    now = datetime.now(timezone.utc).isoformat()
    body_html = result.get("body_html", f"<p>{result.get('body', req.context)}</p>")
    subject = result.get("subject", req.subject)

    with get_db() as db:
        db.execute(
            "INSERT INTO scheduled_emails (id, user_id, tenant_id, case_id, to_email, subject, body_html, status, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, 'draft', ?)",
            (email_id, user["sub"], user["tenant_id"], req.case_id, req.to_email, subject, body_html, now)
        )

    return {
        "email_id": email_id,
        "to": req.to_email,
        "subject": subject,
        "body_html": body_html,
        "status": "draft",
    }


@router.post("/email/{email_id}/send")
async def send_drafted_email(email_id: str, user=Depends(get_current_user)):
    """Send a previously drafted email."""
    with get_db() as db:
        email_row = db.execute(
            "SELECT * FROM scheduled_emails WHERE id = ? AND user_id = ?",
            (email_id, user["sub"])
        ).fetchone()
        if not email_row:
            raise HTTPException(404, "Email not found")
        email_dict = dict(email_row)

    if email_dict["status"] == "sent":
        raise HTTPException(400, "Email already sent")

    # Send via SMTP
    from app.utils.email import _send_email
    success = _send_email(email_dict["to_email"], email_dict["subject"], email_dict["body_html"])

    now = datetime.now(timezone.utc).isoformat()
    new_status = "sent" if success else "failed"
    with get_db() as db:
        db.execute(
            "UPDATE scheduled_emails SET status = ?, sent_at = ? WHERE id = ?",
            (new_status, now if success else None, email_id)
        )

    return {"email_id": email_id, "status": new_status, "sent": success}


@router.get("/emails")
async def list_drafted_emails(user=Depends(get_current_user)):
    """List all drafted/sent emails."""
    with get_db() as db:
        emails = db.execute(
            "SELECT * FROM scheduled_emails WHERE user_id = ? ORDER BY created_at DESC LIMIT 50",
            (user["sub"],)
        ).fetchall()
    return [dict(e) for e in emails]


# ═══════════════════════════════════════════════════════════
# 5. DOCUMENT DRAFTING
# ═══════════════════════════════════════════════════════════

@router.post("/document/draft/public")
async def public_draft_document(req: DocumentDraftRequest):
    """Public document drafting — generates one complete document without auth.
    Users get one free draft, then prompted to sign in for more."""
    model = get_model_for_task("legal_brain_public")

    doc_type_labels = {
        "motion_to_compel": "Motion to Compel Discovery Responses",
        "motion_summary_judgment": "Motion for Summary Judgment",
        "motion_dismiss": "Motion to Dismiss",
        "demand_letter": "Settlement Demand Letter",
        "discovery_request": "Discovery Request (Interrogatories & RFPs)",
        "subpoena": "Subpoena Duces Tecum",
        "declaration": "Declaration / Affidavit",
        "opposition": "Opposition Brief",
        "reply": "Reply Brief",
        "stipulation": "Stipulation and Order",
    }
    doc_label = doc_type_labels.get(req.document_type, req.document_type.replace("_", " ").title())

    system_prompt = f"""You are an elite legal document drafting AI powered by GPT-5.4. Generate a **complete, court-ready, professional** {doc_label}.

You are drafting for experienced litigation attorneys who expect the highest quality work product.

## REQUIREMENTS:
- **Full legal document format** with proper caption, headings, numbered paragraphs, and legal citations
- Include **ALL standard sections** required for this document type — do NOT skip any section
- Use placeholder brackets [BRACKETED TEXT] for case-specific details the attorney must fill in (party names, case numbers, dates, etc.)
- Follow **proper court formatting conventions** for the relevant jurisdiction
- Include a **certificate of service** if applicable
- Include **relevant case citations** and statutory references where appropriate
- Include a **signature block** with appropriate formatting
- The document must be **minimum 3,000-5,000 words** for motions and briefs, **2,000-3,000 words** for demand letters
- Write in **formal legal prose** — the document should be immediately usable after filling in placeholders
- Use proper HTML tags for formatting: <h1>, <h2>, <h3>, <p>, <ol>, <ul>, <li>, <strong>, <em>, <blockquote>, <br>
- Include substantive legal argument, not just procedural boilerplate
- Reference specific rules of court, statutes, and case law that support the arguments

Return JSON: {{"title": "...", "content": "..."}}
Content MUST be in clean HTML format using proper HTML tags — never markdown symbols."""

    result = _call_openai_json(system_prompt,
                                f"Document type: {doc_label}\nContext/Instructions: {req.context}",
                                model=model, max_tokens=16000)

    title = result.get("title", doc_label)
    content = result.get("content", "")

    return {
        "document_id": generate_id(),
        "title": title,
        "content": content,
        "document_type": req.document_type,
        "status": "generated",
    }


@router.post("/document/draft")
async def ai_draft_document(req: DocumentDraftRequest, user=Depends(get_current_user)):
    """Generate a full legal document using AI."""
    user_id = user["sub"]
    with get_db() as db:
        model, _doc_cost = credit_gate(user_id, "legal_brain_draft_document", db)
        _doc_status = db.execute("SELECT subscription_status FROM users WHERE id=?", (user_id,)).fetchone()["subscription_status"]

        case_context = ""
        if req.case_id:
            ctx = _get_case_context(db, req.case_id, user["tenant_id"])
            case_context = _format_case_context(ctx)

    doc_type_labels = {
        "motion_to_compel": "Motion to Compel Discovery Responses",
        "motion_summary_judgment": "Motion for Summary Judgment",
        "motion_dismiss": "Motion to Dismiss",
        "demand_letter": "Settlement Demand Letter",
        "discovery_request": "Discovery Request (Interrogatories & RFPs)",
        "subpoena": "Subpoena Duces Tecum",
        "declaration": "Declaration / Affidavit",
        "opposition": "Opposition Brief",
        "reply": "Reply Brief",
        "stipulation": "Stipulation and Order",
    }
    doc_label = doc_type_labels.get(req.document_type, req.document_type.replace("_", " ").title())

    system_prompt = f"""You are an elite legal document drafting AI powered by GPT-5.4. Generate a **complete, court-ready, professional** {doc_label}.

You are drafting for experienced litigation attorneys who expect the highest quality work product.

{f"## Case Context:{chr(10)}{case_context}" if case_context else ""}

## REQUIREMENTS:
- **Full legal document format** with proper caption, headings, numbered paragraphs, and legal citations
- Include **ALL standard sections** required for this document type — do NOT skip any section
- Use placeholder brackets [BRACKETED TEXT] for case-specific details the attorney must fill in (party names, case numbers, dates, etc.)
- Follow **proper court formatting conventions** for the relevant jurisdiction
- Include a **certificate of service** if applicable
- Include **relevant case citations** and statutory references where appropriate
- Include a **signature block** with appropriate formatting
- The document must be **minimum 3,000-5,000 words** for motions and briefs, **2,000-3,000 words** for demand letters
- Write in **formal legal prose** — the document should be immediately usable after filling in placeholders
- Use proper HTML tags for formatting: <h1>, <h2>, <h3>, <p>, <ol>, <ul>, <li>, <strong>, <em>, <blockquote>, <br>
- Include substantive legal argument, not just procedural boilerplate
- Reference specific rules of court, statutes, and case law that support the arguments
- If case context is provided, tailor the document to the specific facts and legal issues

Return JSON: {{"title": "...", "content": "..."}}
Content MUST be in clean HTML format using proper HTML tags — never markdown symbols."""

    result = _call_openai_json(system_prompt,
                                f"Document type: {doc_label}\nContext/Instructions: {req.context}",
                                model=model, max_tokens=16000)

    doc_id = generate_id()
    now = datetime.now(timezone.utc).isoformat()
    title = result.get("title", doc_label)
    content = result.get("content", "")

    with get_db() as db:
        db.execute(
            "INSERT INTO ai_generated_drafts (id, user_id, tenant_id, case_id, document_type, title, content, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (doc_id, user["sub"], user["tenant_id"], req.case_id, req.document_type, title, content, now)
        )
        deduct_credits(user_id, _doc_status, _doc_cost, "legal_brain_draft_document", db)

    return {
        "document_id": doc_id,
        "title": title,
        "content": content,
        "document_type": req.document_type,
        "status": "generated",
    }


@router.get("/documents")
async def list_ai_drafts(user=Depends(get_current_user)):
    """List all AI-generated document drafts."""
    with get_db() as db:
        drafts = db.execute(
            "SELECT * FROM ai_generated_drafts WHERE user_id = ? ORDER BY created_at DESC LIMIT 50",
            (user["sub"],)
        ).fetchall()
    return [dict(d) for d in drafts]


@router.get("/documents/{doc_id}")
async def get_ai_draft(doc_id: str, user=Depends(get_current_user)):
    """Get a specific AI-generated document."""
    with get_db() as db:
        doc = db.execute(
            "SELECT * FROM ai_generated_drafts WHERE id = ? AND user_id = ?",
            (doc_id, user["sub"])
        ).fetchone()
        if not doc:
            raise HTTPException(404, "Document not found")
    return dict(doc)


# ═══════════════════════════════════════════════════════════
# 6. RESEARCH & ANALYSIS
# ═══════════════════════════════════════════════════════════

@router.post("/research")
async def legal_research(req: PublicQuestion, user=Depends(get_current_user)):
    """Deep legal research with case law citations and analysis."""
    user_id = user["sub"]
    with get_db() as _db:
        model, _res_cost = credit_gate(user_id, "legal_brain_research", _db)
        _res_status = _db.execute("SELECT subscription_status FROM users WHERE id=?", (user_id,)).fetchone()["subscription_status"]

    # Search CourtListener for real, verified case law
    courtlistener_results = _search_courtlistener(req.question, max_results=20)
    logger.info(f"[LEGAL BRAIN RESEARCH] CourtListener returned {len(courtlistener_results)} chars for: {req.question[:80]}")

    system_prompt = f"""You are the **LitigationSpace Legal Research Engine** — powered by GPT-5.4 and backed by CourtListener's verified case law database.

You provide **exhaustive, partner-level legal research** with verified case citations.

## CITATION INTEGRITY (NON-NEGOTIABLE)
- **USE the CourtListener cases below** as your PRIMARY citation source — these are verified real cases
- Copy the **EXACT case name, citation, court, and date** from CourtListener results
- Include the **CourtListener URL** for each cited case
- You may also cite well-known landmark cases you are highly confident about
- **NEVER fabricate** case names, citations, or holdings

{f"### VERIFIED CASE LAW FROM COURTLISTENER{chr(10)}{courtlistener_results}" if courtlistener_results else "*No CourtListener results for this query. Only cite cases you are highly confident are real.*"}

## RESEARCH REQUIREMENTS
Use **rich markdown formatting** (## headings, **bold**, *italics*, > blockquotes, numbered lists, bullet points, [links](url)).

1. Cite specific cases with **full citations** and explain the holding thoroughly
2. Reference relevant statutes with section numbers
3. Note the jurisdiction — federal vs. state, which circuit/state
4. Explain the legal standard clearly and in depth
5. Distinguish between binding and persuasive authority
6. Note any circuit splits or evolving areas of law
7. Provide practical application — litigation strategy implications
8. Discuss counterarguments and how to address them
9. Consider the strength of the position and potential weaknesses
10. **Minimum 3,000-5,000 words** for substantive research questions
11. End with a **Practice Tip** section and **Recommended Next Steps**"""

    try:
        response = _call_openai_with_search(system_prompt, req.question, model=model, max_tokens=16000)
    except Exception:
        response = _call_openai(system_prompt, req.question, model=model, max_tokens=16000)

    response += f"\n\n---\n*Research powered by LitigationSpace Legal Brain + CourtListener Verified Case Law*"

    with get_db() as _db:
        deduct_credits(user_id, _res_status, _res_cost, "legal_brain_research", _db)

    return {"response": response}


# ═══════════════════════════════════════════════════════════
# 7. CALENDAR & REMINDERS
# ═══════════════════════════════════════════════════════════

@router.post("/reminders")
async def create_reminder(req: ReminderRequest, user=Depends(get_current_user)):
    """Create a reminder with optional email notification."""
    reminder_id = generate_id()
    now = datetime.now(timezone.utc).isoformat()

    with get_db() as db:
        db.execute(
            "INSERT INTO reminders (id, user_id, tenant_id, case_id, title, notes, remind_at, remind_via, created_at) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (reminder_id, user["sub"], user["tenant_id"], req.case_id, req.title, req.notes or "", req.remind_at, req.remind_via, now)
        )

    return {"reminder_id": reminder_id, "title": req.title, "remind_at": req.remind_at}


@router.get("/reminders")
async def list_reminders(user=Depends(get_current_user)):
    """List all reminders."""
    with get_db() as db:
        reminders = db.execute(
            "SELECT * FROM reminders WHERE user_id = ? ORDER BY remind_at ASC",
            (user["sub"],)
        ).fetchall()
    return [dict(r) for r in reminders]


@router.delete("/reminders/{reminder_id}")
async def delete_reminder(reminder_id: str, user=Depends(get_current_user)):
    """Delete a reminder."""
    with get_db() as db:
        db.execute("DELETE FROM reminders WHERE id = ? AND user_id = ?", (reminder_id, user["sub"]))
    return {"message": "Reminder deleted"}


@router.get("/briefing")
async def daily_briefing(user=Depends(get_current_user)):
    """Generate an AI-powered daily briefing of everything that needs attention."""
    tenant_id = user["tenant_id"]
    now = datetime.now(timezone.utc)
    today_str = now.strftime("%Y-%m-%d")
    week_str = (now + timedelta(days=7)).strftime("%Y-%m-%d")

    with get_db() as db:
        # Overdue tasks
        overdue_tasks = [dict(r) for r in db.execute("""
            SELECT t.title, t.due_date, t.priority, c.title as case_title FROM tasks t
            JOIN cases c ON t.case_id = c.id
            WHERE t.tenant_id = ? AND t.status IN ('pending', 'in_progress')
            AND t.due_date IS NOT NULL AND t.due_date < ?
            ORDER BY t.due_date ASC LIMIT 10
        """, (tenant_id, today_str)).fetchall()]

        # Due today
        today_tasks = [dict(r) for r in db.execute("""
            SELECT t.title, t.due_date, t.priority, c.title as case_title FROM tasks t
            JOIN cases c ON t.case_id = c.id
            WHERE t.tenant_id = ? AND t.status IN ('pending', 'in_progress')
            AND t.due_date LIKE ?
            ORDER BY t.priority DESC
        """, (tenant_id, f"{today_str}%")).fetchall()]

        # Upcoming this week
        upcoming_tasks = [dict(r) for r in db.execute("""
            SELECT t.title, t.due_date, t.priority, c.title as case_title FROM tasks t
            JOIN cases c ON t.case_id = c.id
            WHERE t.tenant_id = ? AND t.status IN ('pending', 'in_progress')
            AND t.due_date > ? AND t.due_date <= ?
            ORDER BY t.due_date ASC LIMIT 10
        """, (tenant_id, today_str, week_str)).fetchall()]

        # Active cases
        active_cases = [dict(r) for r in db.execute(
            "SELECT title, case_type, status, filing_deadline FROM cases WHERE tenant_id = ? AND status = 'active' ORDER BY filing_deadline ASC LIMIT 10",
            (tenant_id,)
        ).fetchall()]

        # Today's reminders
        today_reminders = [dict(r) for r in db.execute(
            "SELECT title, notes, remind_at FROM reminders WHERE user_id = ? AND status = 'pending' AND remind_at LIKE ? ORDER BY remind_at ASC",
            (user["sub"], f"{today_str}%")
        ).fetchall()]

        # Pending discovery
        overdue_discovery = [dict(r) for r in db.execute("""
            SELECT di.item_description, di.date_due, c.title as case_title FROM discovery_items di
            JOIN cases c ON di.case_id = c.id
            WHERE di.tenant_id = ? AND di.status = 'pending'
            AND di.date_due IS NOT NULL AND di.date_due < ?
            ORDER BY di.date_due ASC LIMIT 5
        """, (tenant_id, today_str)).fetchall()]

    # Build briefing data
    briefing_data = {
        "date": now.strftime("%A, %B %d, %Y"),
        "overdue_tasks": overdue_tasks,
        "today_tasks": today_tasks,
        "upcoming_tasks": upcoming_tasks,
        "active_cases": active_cases,
        "today_reminders": today_reminders,
        "overdue_discovery": overdue_discovery,
    }

    # Generate AI summary
    context = json.dumps(briefing_data, indent=2, default=str)
    system_prompt = """You are a legal briefing AI. Generate a concise daily briefing for a litigation attorney.
Format with clear sections, priorities, and action items. Use markdown. Be direct and actionable.
Start with the most urgent items. Include a "Today's Top 3 Priorities" section."""

    ai_summary = _call_openai(system_prompt, f"Daily briefing data:\n{context}", max_tokens=1500)

    return {
        "briefing": briefing_data,
        "ai_summary": ai_summary,
    }


# ═══════════════════════════════════════════════════════════
# 8. SMART PROACTIVE SUGGESTIONS
# ═══════════════════════════════════════════════════════════

@router.get("/suggestions")
async def get_smart_suggestions(case_id: Optional[str] = None, user=Depends(get_current_user)):
    """AI-powered proactive suggestions for the user's cases."""
    tenant_id = user["tenant_id"]

    with get_db() as db:
        if case_id:
            ctx = _get_case_context(db, case_id, tenant_id)
            cases_data = [ctx] if ctx.get("case") else []
        else:
            cases = db.execute(
                "SELECT id FROM cases WHERE tenant_id = ? AND status = 'active' ORDER BY updated_at DESC LIMIT 5",
                (tenant_id,)
            ).fetchall()
            cases_data = [_get_case_context(db, dict(c)["id"], tenant_id) for c in cases]

    if not cases_data:
        return {"suggestions": [{"type": "info", "message": "No active cases. Create a case to get started!", "priority": "low"}]}

    # Build context for AI
    context_parts = []
    for ctx in cases_data:
        if ctx.get("case"):
            context_parts.append(_format_case_context(ctx))

    full_context = "\n\n---\n\n".join(context_parts)

    system_prompt = """You are a proactive legal assistant AI. Analyze the user's case data and identify:
1. Missing documents that should be uploaded for this case type
2. Approaching deadlines that need attention
3. Tasks that are stalled or blocked
4. Discovery items that are overdue
5. Strategic opportunities (e.g., "consider filing a motion to compel since discovery is overdue")
6. Witness preparation needs
7. Filing requirements that may be missed

Return JSON:
{
  "suggestions": [
    {
      "type": "warning|action|info|deadline",
      "category": "documents|tasks|discovery|strategy|filing|witness",
      "case_title": "...",
      "message": "...",
      "priority": "high|medium|low",
      "action_label": "...",
      "action_type": "create_task|upload_doc|draft_motion|set_reminder|null"
    }
  ]
}

Limit to the 10 most important suggestions, sorted by priority."""

    result = _call_openai_json(system_prompt, f"Case data:\n{full_context}", max_tokens=2000)
    suggestions = result.get("suggestions", [])

    return {"suggestions": suggestions}


# ═══════════════════════════════════════════════════════════
# CONVERSATION HISTORY
# ═══════════════════════════════════════════════════════════

@router.get("/conversations")
async def list_conversations(user=Depends(get_current_user)):
    """List user's Legal Brain conversations with message preview."""
    with get_db() as db:
        convos = db.execute("""
            SELECT lbc.*, c.title as case_title,
            (SELECT COUNT(*) FROM legal_brain_messages WHERE conversation_id = lbc.id) as message_count,
            (SELECT content FROM legal_brain_messages WHERE conversation_id = lbc.id AND role = 'user' ORDER BY created_at ASC LIMIT 1) as preview
            FROM legal_brain_conversations lbc
            LEFT JOIN cases c ON lbc.case_id = c.id
            WHERE lbc.user_id = ?
            ORDER BY lbc.updated_at DESC LIMIT 50
        """, (user["sub"],)).fetchall()
    return [dict(c) for c in convos]


@router.get("/conversations/{conv_id}")
async def get_conversation(conv_id: str, user=Depends(get_current_user)):
    """Get messages in a conversation."""
    with get_db() as db:
        conv = db.execute(
            "SELECT * FROM legal_brain_conversations WHERE id = ? AND user_id = ?",
            (conv_id, user["sub"])
        ).fetchone()
        if not conv:
            raise HTTPException(404, "Conversation not found")
        messages = db.execute(
            "SELECT * FROM legal_brain_messages WHERE conversation_id = ? ORDER BY created_at ASC",
            (conv_id,)
        ).fetchall()
    return {"conversation": dict(conv), "messages": [dict(m) for m in messages]}


# ═══════════════════════════════════════════════════════════
# CRON: Send due reminders via email
# ═══════════════════════════════════════════════════════════

async def send_due_reminders():
    """Check for due reminders and send emails. Called by cron."""
    now = datetime.now(timezone.utc).isoformat()
    from app.utils.email import _send_email

    with get_db() as db:
        due_reminders = db.execute(
            "SELECT r.*, u.email, u.full_name FROM reminders r JOIN users u ON r.user_id = u.id WHERE r.status = 'pending' AND r.remind_at <= ?",
            (now,)
        ).fetchall()

        for reminder in due_reminders:
            r = dict(reminder)
            subject = f"⏰ Reminder: {r['title']}"
            html = f"""
            <div style="font-family: -apple-system, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
                <h2 style="color: #1e293b;">⏰ Reminder</h2>
                <div style="background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 8px; padding: 20px;">
                    <h3 style="margin-top:0;">{r['title']}</h3>
                    {f"<p>{r['notes']}</p>" if r.get('notes') else ""}
                    <p style="color: #64748b; font-size: 13px;">Scheduled for: {r['remind_at']}</p>
                </div>
                <p style="color: #94a3b8; font-size: 12px; margin-top: 20px;">— LitigationSpace Legal Brain</p>
            </div>
            """
            _send_email(r["email"], subject, html)
            db.execute("UPDATE reminders SET status = 'sent' WHERE id = ?", (r["id"],))

    return {"processed": len(due_reminders) if due_reminders else 0}


async def send_daily_briefing_emails():
    """Send daily briefing emails to all users. Called by cron at 7am."""
    from app.utils.email import _send_email

    with get_db() as db:
        users = db.execute("SELECT id, email, full_name, tenant_id FROM users WHERE email_verified = 1").fetchall()

    for user_row in users:
        u = dict(user_row)
        try:
            # Build minimal briefing
            tenant_id = u["tenant_id"]
            today_str = datetime.now(timezone.utc).strftime("%Y-%m-%d")
            with get_db() as db:
                overdue_count = db.execute(
                    "SELECT COUNT(*) as cnt FROM tasks WHERE tenant_id = ? AND status IN ('pending','in_progress') AND due_date IS NOT NULL AND due_date < ?",
                    (tenant_id, today_str)
                ).fetchone()["cnt"]
                today_count = db.execute(
                    "SELECT COUNT(*) as cnt FROM tasks WHERE tenant_id = ? AND status IN ('pending','in_progress') AND due_date LIKE ?",
                    (tenant_id, f"{today_str}%")
                ).fetchone()["cnt"]
                active_count = db.execute(
                    "SELECT COUNT(*) as cnt FROM cases WHERE tenant_id = ? AND status = 'active'",
                    (tenant_id,)
                ).fetchone()["cnt"]

            if overdue_count == 0 and today_count == 0:
                continue  # No briefing needed

            subject = f"📋 Daily Briefing — {overdue_count} overdue, {today_count} due today"
            html = f"""
            <div style="font-family: -apple-system, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px;">
                <h2 style="color: #1e293b;">📋 Your Daily Legal Briefing</h2>
                <p>Good morning, {u['full_name']}!</p>
                <div style="background: #f8fafc; border: 1px solid #e2e8f0; border-radius: 8px; padding: 20px;">
                    <p><strong>🔴 Overdue:</strong> {overdue_count} tasks</p>
                    <p><strong>📅 Due Today:</strong> {today_count} tasks</p>
                    <p><strong>📂 Active Cases:</strong> {active_count}</p>
                </div>
                <div style="text-align:center; margin: 20px 0;">
                    <a href="https://litigationspace.com/dashboard" style="background: linear-gradient(135deg, #7c3aed, #6366f1); color: white; padding: 12px 24px; border-radius: 8px; text-decoration: none; font-weight: 600;">Open Dashboard</a>
                </div>
                <p style="color: #94a3b8; font-size: 12px;">— LitigationSpace Legal Brain</p>
            </div>
            """
            _send_email(u["email"], subject, html)
        except Exception as e:
            logger.error(f"[BRIEFING] Failed for {u['email']}: {e}")


# -- Document Analyzer Endpoints ----------------------------------------------
# POST /analyze-documents, /analysis-followup, /download

from fastapi import UploadFile, File, Form
from fastapi.responses import Response as FastAPIResponse
from typing import List
import fitz  # pymupdf
import io as _io
import uuid as _uuid


def _da_extract_text(file_bytes: bytes, filename: str, content_type: str = "") -> str:
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    try:
        if ext == "pdf" or "pdf" in content_type:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return "\n\n".join(page.get_text() for page in doc)[:80000]
        elif ext in ("docx", "doc"):
            from docx import Document as DocxDocument
            doc = DocxDocument(_io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:80000]
        else:
            return file_bytes.decode("utf-8", errors="replace")[:80000]
    except Exception as e:
        logger.warning("[DA] text extraction failed for %s: %s", filename, e)
        return ""


@router.post("/analyze-documents")
async def analyze_documents(
    files: List[UploadFile] = File(...),
    analysis_type: str = Form("comprehensive"),
    instruction: str = Form("Analyze this document comprehensively."),
    document_type: str = Form("Legal Document"),
    current_user: dict = Depends(get_current_user),
):
    user_id = current_user["sub"]

    with get_db() as db:
        ai_model, credit_cost = credit_gate(user_id, "document_analysis", db)
        row = db.execute("SELECT subscription_status FROM users WHERE id = ?", (user_id,)).fetchone()
        sub_status = row["subscription_status"] if row else "trial"

    texts = []
    doc_infos = []
    for f in files:
        raw = await f.read()
        text = _da_extract_text(raw, f.filename or "file.txt", f.content_type or "")
        if text.strip():
            header = "=== " + (f.filename or "file.txt") + " ==="
            texts.append(header + "\n" + text)
            doc_infos.append({
                "filename": f.filename or "file.txt",
                "size_kb": round(len(raw) / 1024, 1),
                "text_length": len(text),
            })

    if not texts:
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from the uploaded document(s). "
                   "Please ensure the file is not password-protected or empty.",
        )

    combined = "\n\n".join(texts)[:100000]

    system_prompt = (
        "You are an expert legal document analyst specializing in " + document_type + ". "
        "Provide thorough, structured, actionable analysis in markdown format. "
        "Use clear headings (##), bullet points, and highlight critical clauses, "
        "risks, and recommendations."
    )

    analysis_prompts = {
        "comprehensive": (
            "Perform a comprehensive analysis of this " + document_type + ". "
            "Cover: key terms and obligations, risks, unusual or missing clauses, "
            "recommendations, and overall assessment."
        ),
        "contract_review": (
            "Review this " + document_type + " for legal soundness. "
            "Identify unfavorable terms, missing protections, ambiguous language, "
            "and negotiation opportunities."
        ),
        "risk_assessment": (
            "Identify all legal, financial, and operational risks in this "
            + document_type + ". Rate each risk (High/Medium/Low) and suggest mitigations."
        ),
        "clause_extraction": (
            "Extract and explain every significant clause in this " + document_type + ". "
            "For each clause: title, plain-language explanation, and implications."
        ),
        "summary": (
            "Write an executive summary of this " + document_type + " for a non-lawyer. "
            "Bullet the key points, obligations, deadlines, and red flags."
        ),
        "comparison": (
            "Analyze this " + document_type + " against standard industry benchmarks. "
            "Note deviations, advantages, and disadvantages."
        ),
        "redline_suggestions": (
            "Suggest specific redline changes to this " + document_type + ". "
            "For each suggestion: original text, recommended revision, and reason."
        ),
        "case_law_extraction": (
            "Identify all legal citations, statutes, and case law referenced in this "
            "document. Summarize each and assess their relevance."
        ),
        "discovery_review": (
            "Review this document for discovery relevance. Categorize by relevance level "
            "and identify key facts, admissions, and potential evidence."
        ),
        "custom": instruction,
    }
    user_msg = analysis_prompts.get(analysis_type, instruction)

    client = _get_openai_client()
    conversation_id = str(_uuid.uuid4())

    try:
        resp = client.chat.completions.create(
            model=ai_model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_msg + "\n\n---\n\n" + combined},
            ],
            max_completion_tokens=16000,
        )
        analysis_text = resp.choices[0].message.content or ""
    except Exception as e:
        logger.error("[DA] OpenAI error: %s", e)
        raise HTTPException(status_code=502, detail="AI analysis failed. Please try again.")

    if not analysis_text.strip():
        logger.error("[DA] empty analysis_text, finish_reason=%s", getattr(resp.choices[0], "finish_reason", "?"))
        raise HTTPException(
            status_code=502,
            detail="The AI returned an empty analysis, likely because the document(s) are too large. "
                   "Try analyzing fewer or shorter documents.",
        )

    with get_db() as db:
        deduct_credits(user_id, sub_status, credit_cost, "document_analysis", db)
        try:
            db.execute(
                "INSERT INTO document_analysis_conversations "
                "(id, user_id, document_type, analysis_type, document_text, analysis_text, created_at) "
                "VALUES (?, ?, ?, ?, ?, ?, ?)",
                (
                    conversation_id, user_id, document_type, analysis_type,
                    combined[:50000], analysis_text,
                    datetime.now(timezone.utc).isoformat(),
                ),
            )
        except Exception:
            pass  # table may not exist — non-fatal

    return {
        "conversation_id": conversation_id,
        "response": analysis_text,
        "model": ai_model,
        "files_processed": len(doc_infos),
        "documents": doc_infos,
    }


@router.post("/analysis-followup")
async def analysis_followup(
    req: dict,
    current_user: dict = Depends(get_current_user),
):
    user_id = current_user["sub"]
    conversation_id = req.get("conversation_id", "")
    action = req.get("action", "")
    party_role = req.get("party_role", "")
    jurisdiction = req.get("jurisdiction", "")
    custom_instruction = req.get("custom_instruction", "")

    prior_doc = prior_analysis = ""
    doc_type = "Legal Document"
    with get_db() as db:
        try:
            row = db.execute(
                "SELECT document_text, analysis_text, document_type "
                "FROM document_analysis_conversations WHERE id = ? AND user_id = ?",
                (conversation_id, user_id),
            ).fetchone()
            if row:
                prior_doc = row["document_text"] or ""
                prior_analysis = row["analysis_text"] or ""
                doc_type = row["document_type"] or "Legal Document"
        except Exception:
            pass

    role_str = party_role or "responding party"
    jur_str = jurisdiction or "applicable"
    followup_prompts = {
        "draft_response": (
            "Draft a formal legal response to this " + doc_type
            + " on behalf of the " + role_str + " under " + jur_str + " law."
        ),
        "draft_counter": (
            "Draft a counter-proposal to this " + doc_type
            + " with improved terms for the " + role_str + " under " + jur_str + " law."
        ),
        "negotiation_points": (
            "List the top 10 negotiation points in this " + doc_type
            + " for the " + role_str
            + ", with suggested positions and fallback positions."
        ),
        "compliance_check": (
            "Check this " + doc_type + " for compliance with " + jur_str
            + " law and regulations. List any violations or gaps."
        ),
        "explain_simple": (
            "Explain this " + doc_type + " in plain language a non-lawyer can understand. "
            "Avoid jargon."
        ),
        "identify_loopholes": (
            "Identify any loopholes, ambiguities, or vulnerabilities in this "
            + doc_type + " that could be exploited."
        ),
        "custom": custom_instruction or "Provide additional analysis.",
    }
    prompt = followup_prompts.get(action, custom_instruction or "Provide additional analysis.")

    if prior_analysis:
        context = "Original document analysis:\n" + prior_analysis + "\n\nDocument text:\n" + prior_doc[:30000]
    else:
        context = prior_doc[:50000]

    with get_db() as db:
        ai_model, credit_cost = credit_gate(user_id, "document_analysis", db)
        row = db.execute("SELECT subscription_status FROM users WHERE id = ?", (user_id,)).fetchone()
        sub_status = row["subscription_status"] if row else "trial"

    client = _get_openai_client()
    try:
        msgs = [{"role": "system", "content": "You are an expert legal analyst. Respond in markdown format."}]
        if context:
            msgs.append({"role": "assistant", "content": context})
        msgs.append({"role": "user", "content": prompt})
        resp = client.chat.completions.create(
            model=ai_model, messages=msgs, max_completion_tokens=3000
        )
        followup_text = resp.choices[0].message.content or ""
    except Exception as e:
        logger.error("[DA-FOLLOWUP] OpenAI error: %s", e)
        raise HTTPException(status_code=502, detail="Follow-up generation failed.")

    with get_db() as db:
        deduct_credits(user_id, sub_status, credit_cost, "document_analysis", db)

    return {"response": followup_text}


@router.post("/download")
async def download_analysis(
    req: dict,
    current_user: dict = Depends(get_current_user),
):
    content = req.get("content", "")
    title = req.get("title", "Document Analysis")
    fmt = req.get("format", "docx")
    safe_title = title.replace(" ", "_")[:40]

    if fmt == "docx":
        from docx import Document as DocxDocument
        from docx.shared import RGBColor
        NAVY = RGBColor(0x0A, 0x0F, 0x1E)
        GOLD = RGBColor(0xC8, 0x99, 0x2A)
        doc = DocxDocument()
        doc.core_properties.title = title
        h = doc.add_heading(title, 0)
        if h.runs:
            h.runs[0].font.color.rgb = NAVY
        for line in content.split("\n"):
            line = line.rstrip()
            if line.startswith("### "):
                hh = doc.add_heading(line[4:], level=3)
                if hh.runs:
                    hh.runs[0].font.color.rgb = GOLD
            elif line.startswith("## "):
                hh = doc.add_heading(line[3:], level=2)
                if hh.runs:
                    hh.runs[0].font.color.rgb = NAVY
            elif line.startswith("# "):
                hh = doc.add_heading(line[2:], level=1)
                if hh.runs:
                    hh.runs[0].font.color.rgb = NAVY
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.strip() == "---":
                doc.add_paragraph("_" * 60)
            elif line.strip():
                doc.add_paragraph(line)
        buf = _io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return FastAPIResponse(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": 'attachment; filename="' + safe_title + '.docx"'},
        )
    else:
        from fpdf import FPDF
        FONT_DIR = "/usr/share/fonts/truetype/liberation"
        NAVY = (10, 15, 30)
        GOLD = (200, 153, 42)
        BODY = (40, 40, 40)

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_font("Liberation", "", f"{FONT_DIR}/LiberationSans-Regular.ttf")
        pdf.add_font("Liberation", "B", f"{FONT_DIR}/LiberationSans-Bold.ttf")
        pdf.add_font("Liberation", "I", f"{FONT_DIR}/LiberationSans-Italic.ttf")

        def cell(text, h=6):
            pdf.multi_cell(0, h, text)
            pdf.set_x(pdf.l_margin)

        pdf.set_text_color(*NAVY)
        pdf.set_font("Liberation", "B", 18)
        cell(title[:80], 12)
        pdf.set_draw_color(*GOLD)
        pdf.set_line_width(0.8)
        pdf.line(10, pdf.get_y() + 1, 200, pdf.get_y() + 1)
        pdf.ln(6)

        pdf.set_text_color(*BODY)
        pdf.set_font("Liberation", "", 11)
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                pdf.ln(3)
            elif line.startswith("### "):
                pdf.set_text_color(*GOLD)
                pdf.set_font("Liberation", "B", 12)
                cell(line[4:].strip(), 7)
                pdf.set_text_color(*BODY)
                pdf.set_font("Liberation", "", 11)
            elif line.startswith("## ") or line.startswith("# "):
                pdf.set_text_color(*NAVY)
                pdf.set_font("Liberation", "B", 14)
                cell(line.lstrip("#").strip(), 8)
                pdf.set_text_color(*BODY)
                pdf.set_font("Liberation", "", 11)
                pdf.ln(1)
            elif line.startswith("- ") or line.startswith("* "):
                cell("•  " + line[2:])
            elif line == "---":
                pdf.set_draw_color(*GOLD)
                pdf.set_line_width(0.4)
                pdf.line(10, pdf.get_y() + 2, 200, pdf.get_y() + 2)
                pdf.ln(6)
            else:
                cell(line)
        pdf_bytes = pdf.output()
        return FastAPIResponse(
            content=bytes(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": 'attachment; filename="' + safe_title + '.pdf"'},
        )
"""
RAG Litigation Engine — management API endpoints.
Appended to legal_brain.py router.
"""


@router.get("/rag/status")
async def rag_status(current_user: dict = Depends(get_current_user)):
    """Return current RAG index coverage statistics."""
    try:
        from app.utils.rag_engine import get_index_status
        with get_db() as db:
            status = get_index_status(db)
        return {"ok": True, **status}
    except Exception as e:
        raise HTTPException(500, f"RAG status error: {e}")


@router.post("/rag/reindex")
async def rag_reindex(req: dict = {}, current_user: dict = Depends(get_current_user)):
    """
    Trigger background re-indexing of unembedded chunks.
    Pass {"limit": N} to index only N chunks (default: all unembedded).
    Pass {"force": true} to re-embed everything.
    """
    import threading

    limit = req.get("limit") if req else None
    force = bool(req.get("force", False)) if req else False

    def _run():
        try:
            from app.utils.rag_engine import index_chunks, invalidate_cache
            with get_db() as db:
                result = index_chunks(db, limit=limit, force=force)
            invalidate_cache()  # clear cache so next search reloads
            logger.info(f"[RAG REINDEX] Complete: {result}")
        except Exception as e:
            logger.error(f"[RAG REINDEX] Failed: {e}")

    t = threading.Thread(target=_run, daemon=True)
    t.start()

    return {
        "ok":     True,
        "status": "indexing_started",
        "message": f"Background indexing started (limit={limit}, force={force}). Check /api/legal-brain/rag/status for progress.",
    }


@router.post("/rag/index-document/{doc_id}")
async def rag_index_document(doc_id: str, current_user: dict = Depends(get_current_user)):
    """Index a specific document by ID."""
    try:
        from app.utils.rag_engine import index_document
        with get_db() as db:
            count = index_document(db, doc_id)
        return {"ok": True, "chunks_indexed": count, "document_id": doc_id}
    except Exception as e:
        raise HTTPException(500, f"Index error: {e}")


@router.post("/rag/search-test")
async def rag_search_test(req: dict, current_user: dict = Depends(get_current_user)):
    """
    Test the RAG search directly — useful for debugging.
    Body: {"question": "...", "jurisdiction": "US" (optional)}
    """
    question    = req.get("question", "")
    jurisdiction = req.get("jurisdiction")
    if not question:
        raise HTTPException(400, "question is required")
    try:
        from app.utils.rag_engine import search, _ensure_embedding_column
        with get_db() as db:
            _ensure_embedding_column(db)
            results = search(db, question, top_k=5, jurisdiction=jurisdiction)
        return {
            "question": question,
            "result_count": len(results),
            "results": results,
        }
    except Exception as e:
        raise HTTPException(500, f"Search error: {e}")


# ── Email Drafter — Send ──────────────────────────────────────────────────────

class SendEmailReq(BaseModel):
    from_email: str
    to_email:   str
    subject:    str
    body:       str

@router.post("/send-email")
def send_drafted_email(req: SendEmailReq):
    """Send a drafted legal email to the recipient and a copy to the sender."""
    from app.utils.email import _send_email
    import re

    email_re = re.compile(r'^[^\s@]+@[^\s@]+\.[^\s@]+$')
    if not email_re.match(req.to_email.strip()):
        raise HTTPException(400, "Invalid recipient email address")
    if not req.subject.strip():
        raise HTTPException(400, "Subject is required")
    if not req.body.strip():
        raise HTTPException(400, "Email body is required")

    def to_html(text):
        t = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        return '<br>'.join(t.splitlines())

    body_html = to_html(req.body)

    main_html = (
        '<div style="font-family:Georgia,serif;font-size:14px;line-height:1.8;color:#111;max-width:680px;margin:0 auto;padding:24px;">' +
        '<div style="border-left:4px solid #D4950E;padding-left:16px;margin-bottom:20px;">' +
        '<p style="margin:0;font-size:11px;color:#888;text-transform:uppercase;letter-spacing:.06em;">Sent via LegalBrain &middot; LitigationSpace</p>' +
        '</div>' +
        body_html +
        '<hr style="border:none;border-top:1px solid #ddd;margin:24px 0;">' +
        '<p style="font-size:11px;color:#aaa;margin:0;">Drafted and sent via <a href="https://litigationspace.com" style="color:#C89820;">LitigationSpace LegalBrain</a>.</p>' +
        '</div>'
    )

    errors = []
    ok = _send_email(req.to_email.strip(), req.subject.strip(), main_html)
    if not ok:
        errors.append(f"Failed to deliver to {req.to_email}")

    if req.from_email.strip() and req.from_email.strip().lower() != req.to_email.strip().lower():
        if email_re.match(req.from_email.strip()):
            copy_html = (
                '<div style="font-family:Georgia,serif;font-size:14px;color:#111;max-width:680px;margin:0 auto;padding:24px;">' +
                '<div style="background:#fff8e7;border:1px solid #D4950E;border-radius:8px;padding:12px 16px;margin-bottom:20px;">' +
                '<p style="margin:0;font-size:12px;color:#7a5a00;font-weight:600;">Copy of email you sent via LegalBrain</p>' +
                '</div>' +
                main_html +
                '</div>'
            )
            _send_email(req.from_email.strip(), f"[Copy] {req.subject.strip()}", copy_html)

    if errors:
        raise HTTPException(500, "; ".join(errors))

    return {"status": "sent", "to": req.to_email, "copy_sent_to": req.from_email}


# ── Chat with File Attachments ────────────────────────────────────────────────

def _extract_file_text(raw: bytes, filename: str, content_type: str = "") -> str:
    """Extract text from any document type. Falls back to OCR for images/scanned PDFs."""
    import io as _io2
    ext = (filename.rsplit(".", 1)[-1] if "." in filename else "").lower()
    ct  = (content_type or "").lower()

    try:
        # PDF
        if ext == "pdf" or "pdf" in ct:
            doc = fitz.open(stream=raw, filetype="pdf")
            pages_text = [page.get_text() for page in doc]
            full = "\n\n".join(pages_text).strip()
            # If PDF has no text layer (scanned), run OCR page by page
            if not full:
                from PIL import Image
                import pytesseract
                ocr_parts = []
                for page in doc:
                    pix = page.get_pixmap(dpi=200)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_parts.append(pytesseract.image_to_string(img))
                full = "\n\n".join(ocr_parts).strip()
            return full[:100000]

        # Word DOCX
        elif ext in ("docx",):
            from docx import Document as DocxDoc
            doc = DocxDoc(_io2.BytesIO(raw))
            lines = []
            for para in doc.paragraphs:
                if para.text.strip():
                    lines.append(para.text)
            # Also grab table content
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            return "\n".join(lines)[:100000]

        # Word DOC (older format - read as binary and extract readable text)
        elif ext == "doc":
            # Try docx first (some .doc are actually docx)
            try:
                from docx import Document as DocxDoc
                doc = DocxDoc(_io2.BytesIO(raw))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())[:100000]
            except Exception:
                # Strip binary, extract ASCII runs
                text = raw.decode("latin-1", errors="replace")
                import re as _re3
                words = _re3.findall(r'[A-Za-z0-9][^\x00-\x1f]{3,}', text)
                return " ".join(words)[:50000]

        # Excel XLSX / XLS
        elif ext in ("xlsx", "xls"):
            import openpyxl
            wb = openpyxl.load_workbook(_io2.BytesIO(raw), read_only=True, data_only=True)
            rows_out = []
            for sheet in wb.worksheets:
                rows_out.append(f"[Sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True, max_row=1000):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        rows_out.append("\t".join(cells))
            return "\n".join(rows_out)[:80000]

        # CSV
        elif ext == "csv":
            return raw.decode("utf-8", errors="replace")[:80000]

        # Images — OCR
        elif ext in ("jpg", "jpeg", "png", "tiff", "tif", "bmp", "gif", "webp") or ct.startswith("image/"):
            from PIL import Image
            import pytesseract
            img = Image.open(_io2.BytesIO(raw))
            return pytesseract.image_to_string(img)[:80000]

        # Plain text / everything else
        else:
            try:
                return raw.decode("utf-8", errors="replace")[:80000]
            except Exception:
                return ""

    except Exception as e:
        logger.warning("[FILES] extraction failed for %s: %s", filename, e)
        return ""


IMAGE_EXTS = {"jpg", "jpeg", "png", "gif", "webp", "tiff", "tif", "bmp"}


@router.post("/chat-with-files")
async def chat_with_files(
    files: List[UploadFile] = File(default=[]),
    message: str = Form(default=""),
    history: str = Form(default="[]"),
    credentials: HTTPAuthorizationCredentials = Depends(HTTPBearer(auto_error=False)),
):
    """
    Chat endpoint that accepts up to 20 file attachments.
    Supports PDF, DOCX, DOC, XLSX, XLS, CSV, images, scanned documents, and plain text.
    Scanned PDFs and images are processed with OCR. Images can also be sent to vision model.
    """
    import json, base64

    # Optional auth
    user_id = None
    if credentials:
        try:
            payload = decode_token(credentials.credentials)
            user_id = payload.get("sub")
        except Exception:
            pass

    try:
        chat_history = json.loads(history)
    except Exception:
        chat_history = []

    if not message.strip() and not files:
        raise HTTPException(400, "Provide a message or attach files.")

    # ── Process files ───────────────────────────────────────────────────────
    text_parts    = []   # extracted text sections
    vision_images = []   # base64-encoded images for GPT vision
    file_count    = 0

    for f in (files or [])[:20]:
        raw   = await f.read()
        fname = (f.filename or "file").strip()
        ct    = (f.content_type or "").lower()
        ext   = (fname.rsplit(".", 1)[-1] if "." in fname else "").lower()
        file_count += 1

        if ext in IMAGE_EXTS or ct.startswith("image/"):
            # Send image to vision model AND run OCR as fallback text
            mime = f.content_type or f"image/{ext or 'jpeg'}"
            b64  = base64.b64encode(raw).decode()
            vision_images.append({"name": fname, "mime": mime, "b64": b64})
            # Also try OCR so we have text version
            ocr_text = _extract_file_text(raw, fname, ct)
            if ocr_text.strip():
                text_parts.append(f"=== {fname} (OCR) ===\n{ocr_text}")
        else:
            extracted = _extract_file_text(raw, fname, ct)
            if extracted.strip():
                text_parts.append(f"=== {fname} ===\n{extracted}")
            else:
                text_parts.append(f"=== {fname} ===\n[No text could be extracted from this file]")

    # ── Build AI messages ───────────────────────────────────────────────────
    system_prompt = (
        "You are LegalBrain, an expert AI legal assistant with deep knowledge of law across all jurisdictions. "
        "You can read, analyze, and reason about legal documents, contracts, court filings, evidence, financial records, "
        "and any other document type. When documents are provided:\n"
        "- Read and understand every section carefully\n"
        "- Follow the user's specific instructions about what to do with the documents\n"
        "- Provide thorough, structured, actionable analysis\n"
        "- Identify legal issues, risks, obligations, and strategic opportunities\n"
        "- Use clear headings (##), numbered lists, and bold for key points\n"
        "- Be precise — cite specific clauses, dates, names, and figures from the documents\n"
        "You have REAL-TIME WEB SEARCH capability. Never say 'my knowledge cuts off at'."
    )

    # Determine if we use vision model
    use_vision = len(vision_images) > 0
    model = "gpt-4o"   # gpt-4o supports vision + large context

    # Build user message content
    user_content: list = []

    # Inject document texts
    if text_parts:
        combined_docs = "\n\n---\n\n".join(text_parts)
        total_chars   = len(combined_docs)
        if total_chars > 200000:
            combined_docs = combined_docs[:200000] + "\n\n[... document truncated at 200,000 chars ...]"
        user_content.append({
            "type": "text",
            "text": f"ATTACHED DOCUMENTS ({len(text_parts)} file(s), {total_chars:,} chars extracted):\n\n{combined_docs}"
        })

    # Inject vision images
    for img in vision_images:
        user_content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{img['mime']};base64,{img['b64']}", "detail": "high"}
        })
        user_content.append({"type": "text", "text": f"[Image above: {img['name']}]"})

    # User instruction
    instruction = message.strip() or (
        f"Please analyze the attached document{'s' if file_count > 1 else ''} "
        "in depth and provide a comprehensive assessment."
    )
    user_content.append({"type": "text", "text": instruction})

    # Build messages
    messages_list = [{"role": "system", "content": system_prompt}]
    for h in (chat_history or [])[-8:]:
        role    = h.get("role", "user")
        content = h.get("content", "")
        if role in ("user", "assistant") and content:
            messages_list.append({"role": role, "content": content})

    # Final user message — use list content for vision, string for text-only
    if use_vision or text_parts:
        messages_list.append({"role": "user", "content": user_content})
    else:
        messages_list.append({"role": "user", "content": instruction})

    # Add web search context if query seems to need current info
    search_ctx = _serpapi_search(instruction) if instruction else ""
    if search_ctx and not text_parts and not vision_images:
        messages_list[0]["content"] += (
            "\n\n---\n" + search_ctx +
            "\nUSE the search results above to answer with current, accurate information.\n---"
        )

    # ── Call OpenAI ─────────────────────────────────────────────────────────
    client = _get_openai_client()
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages_list,
            max_completion_tokens=4096,
            temperature=0.2,
        )
        answer = (resp.choices[0].message.content or "").strip()
    except Exception as e:
        logger.error("[CHAT-FILES] OpenAI error: %s", e)
        raise HTTPException(500, f"AI processing error: {e}")

    return {
        "answer": answer,
        "files_processed": file_count,
    }
