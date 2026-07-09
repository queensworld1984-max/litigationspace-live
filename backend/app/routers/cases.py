"""Cases router - CRUD operations with multi-tenant isolation."""
import json
import io
import logging
import os
import shutil
import traceback
from datetime import datetime, timezone
from pathlib import Path
from typing import List, Optional
from pydantic import BaseModel
from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, BackgroundTasks, Query
from fastapi.responses import FileResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials

from app.database import get_db
from app.models.schemas import CaseCreate, CaseUpdate, TaskCreate, TaskUpdate, DocumentCreate
from app.utils.auth import get_current_user, generate_id, decode_token

_bearer = HTTPBearer(auto_error=False)

def get_download_user(
    credentials: HTTPAuthorizationCredentials = Depends(_bearer),
    token: Optional[str] = Query(None),
) -> dict:
    """Auth for download endpoints: accepts Authorization header OR ?token= query param."""
    if credentials:
        return decode_token(credentials.credentials)
    if token:
        return decode_token(token)
    raise HTTPException(status_code=401, detail="Not authenticated")
from app.utils.model_router import get_model_for_task
from app.utils.credits import credit_gate, deduct_credits

logger = logging.getLogger(__name__)

# File upload configuration
UPLOAD_BASE_DIR = os.environ.get("UPLOAD_DIR", "/var/www/litigationspace/data/uploads")
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100MB
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".rtf", ".odt", ".png", ".jpg", ".jpeg", ".gif", ".webp", ".xlsx", ".csv", ".xls", ".pptx", ".ppt", ".tiff", ".tif", ".bmp", ".svg", ".heic", ".heif", ".msg", ".eml", ".pages", ".numbers", ".key"}


def _convert_to_pdf_via_libreoffice(fp: Path) -> Optional[Path]:
    """Convert a document to PDF using LibreOffice. Returns path to converted PDF or None."""
    import subprocess
    import tempfile
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            # Copy file to temp dir to avoid permission issues
            tmp_input = Path(tmpdir) / fp.name
            shutil.copy2(str(fp), str(tmp_input))
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, str(tmp_input)],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode != 0:
                logger.warning(f"LibreOffice conversion failed: {result.stderr}")
                return None
            # Find the output PDF
            pdf_name = tmp_input.stem + ".pdf"
            pdf_path = Path(tmpdir) / pdf_name
            if pdf_path.exists():
                # Move to a stable temp location before tmpdir cleanup
                stable_tmp = Path(tempfile.mktemp(suffix=".pdf"))
                shutil.copy2(str(pdf_path), str(stable_tmp))
                return stable_tmp
            return None
    except Exception as e:
        logger.warning(f"LibreOffice conversion error for {fp}: {e}")
        return None


def _add_file_to_pdf(merged_pdf, fp: Path, doc_label: str = ""):
    """Convert any supported file type to PDF pages and append to merged_pdf.
    Supports: PDF, images, DOCX/DOC, TXT, RTF, CSV, XLSX."""
    import fitz
    ext = fp.suffix.lower()
    try:
        if ext == ".pdf":
            src = fitz.open(str(fp))
            merged_pdf.insert_pdf(src)
            src.close()
        elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
            img_doc = fitz.open()
            img_page = img_doc.new_page(width=612, height=792)
            img_rect = fitz.Rect(36, 36, 576, 756)
            img_page.insert_image(img_rect, filename=str(fp))
            merged_pdf.insert_pdf(img_doc)
            img_doc.close()
        elif ext in (".docx", ".doc"):
            # Use LibreOffice to convert to PDF preserving formatting and all pages
            converted_pdf_path = _convert_to_pdf_via_libreoffice(fp)
            if converted_pdf_path and converted_pdf_path.exists():
                src = fitz.open(str(converted_pdf_path))
                merged_pdf.insert_pdf(src)
                src.close()
                try:
                    converted_pdf_path.unlink()  # clean up temp file
                except Exception:
                    pass
            else:
                # Fallback: extract text (loses formatting)
                doc_bytes = fp.read_bytes()
                text = _extract_text_from_docx(doc_bytes)
                if text:
                    _add_text_pages(merged_pdf, text, doc_label)
        elif ext in (".rtf", ".odt"):
            # Use LibreOffice to convert RTF/ODT to PDF preserving formatting
            converted_pdf_path = _convert_to_pdf_via_libreoffice(fp)
            if converted_pdf_path and converted_pdf_path.exists():
                src = fitz.open(str(converted_pdf_path))
                merged_pdf.insert_pdf(src)
                src.close()
                try:
                    converted_pdf_path.unlink()
                except Exception:
                    pass
            else:
                # Fallback: extract as plain text
                try:
                    text = fp.read_text(encoding="utf-8", errors="replace")
                except Exception:
                    text = fp.read_bytes().decode("utf-8", errors="replace")
                if text.strip():
                    _add_text_pages(merged_pdf, text.strip(), doc_label)
        elif ext == ".txt":
            try:
                text = fp.read_text(encoding="utf-8", errors="replace")
            except Exception:
                text = fp.read_bytes().decode("utf-8", errors="replace")
            if text.strip():
                _add_text_pages(merged_pdf, text.strip(), doc_label)
        elif ext == ".csv":
            text = fp.read_text(encoding="utf-8", errors="replace")
            if text.strip():
                _add_text_pages(merged_pdf, text.strip(), doc_label)
        elif ext == ".xlsx":
            try:
                import openpyxl
                wb = openpyxl.load_workbook(str(fp), data_only=True)
                lines = []
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    lines.append(f"--- Sheet: {sheet} ---")
                    for row in ws.iter_rows(values_only=True):
                        lines.append("\t".join(str(c) if c is not None else "" for c in row))
                text = "\n".join(lines)
                if text.strip():
                    _add_text_pages(merged_pdf, text.strip(), doc_label)
            except ImportError:
                text = f"[XLSX file: {fp.name} — openpyxl not available]"
                _add_text_pages(merged_pdf, text, doc_label)
        else:
            # Unsupported — add a placeholder page
            _add_text_pages(merged_pdf, f"[File: {fp.name} — format not directly convertible]", doc_label)
    except Exception as e:
        logger.warning(f"_add_file_to_pdf failed for {fp}: {e}")


def _add_text_pages(merged_pdf, text: str, title: str = ""):
    """Add text content as PDF pages, handling multi-page overflow."""
    import fitz
    # Split text into chunks that fit on a page
    max_chars_per_page = 3000
    chunks = []
    while text:
        chunks.append(text[:max_chars_per_page])
        text = text[max_chars_per_page:]
    for i, chunk in enumerate(chunks):
        page = merged_pdf.new_page(width=612, height=792)
        # Add title on first page
        if title and i == 0:
            title_rect = fitz.Rect(72, 40, 540, 65)
            page.insert_textbox(title_rect, title, fontsize=12, fontname="helv-b" if hasattr(fitz, 'helv-b') else "helv")
            text_rect = fitz.Rect(72, 72, 540, 750)
        else:
            text_rect = fitz.Rect(72, 50, 540, 750)
        page.insert_textbox(text_rect, chunk, fontsize=10, fontname="helv")


def _apply_bates_numbering(merged_pdf, prefix: str, start_num: int = 1):
    """Stamp Bates numbers on every page of a merged PDF."""
    import fitz
    for page_idx in range(merged_pdf.page_count):
        page = merged_pdf[page_idx]
        bates_num = start_num + page_idx
        bates_text = f"{prefix}-{bates_num:06d}"
        # Position: bottom-right corner
        rect = fitz.Rect(440, 760, 590, 780)
        # White background for readability
        page.draw_rect(rect, color=(0.6, 0.6, 0.6), fill=(1, 1, 1), width=0.5)
        page.insert_textbox(
            fitz.Rect(442, 762, 588, 778),
            bates_text,
            fontsize=8,
            fontname="cour",
            color=(0.2, 0.2, 0.2),
            align=2  # right align
        )


# --------------- Exhibit helpers ---------------

def _exhibit_letter(n: int) -> str:
    """Convert 0-based index to exhibit letter: 0->A, 25->Z, 26->AA, 27->AB, etc."""
    if n < 26:
        return chr(65 + n)
    result = ""
    while n >= 0:
        result = chr(65 + (n % 26)) + result
        n = n // 26 - 1
    return result


def _int_to_roman(n: int) -> str:
    """Convert 1-based integer to Roman numeral string."""
    vals = [(1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
            (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
            (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I")]
    result = ""
    for value, numeral in vals:
        while n >= value:
            result += numeral
            n -= value
    return result


# All supported exhibit numbering schemes (15 formats)
EXHIBIT_SCHEMES = {
    "letters", "numbers", "claimant_numbers", "respondent_numbers",
    "annexture_numbers", "annexture_letters", "plaintiff_numbers",
    "defendant_numbers", "joint_numbers", "government_numbers",
    "roman_numerals", "bates_numbered", "tab_numbers",
    "schedule_letters", "appendix_letters",
}


def _make_exhibit_label(count: int, numbering: str) -> str:
    """Return the exhibit label string for count (0-based) and numbering scheme."""
    n = count + 1  # 1-based for display
    if numbering == "numbers":
        return str(n)
    elif numbering == "claimant_numbers":
        return f"C-{n}"
    elif numbering == "respondent_numbers":
        return f"R-{n}"
    elif numbering == "annexture_numbers":
        return f"Annexture {n}"
    elif numbering == "annexture_letters":
        return f"Annexture {_exhibit_letter(count)}"
    elif numbering == "plaintiff_numbers":
        return f"P-{n}"
    elif numbering == "defendant_numbers":
        return f"D-{n}"
    elif numbering == "joint_numbers":
        return f"J-{n}"
    elif numbering == "government_numbers":
        return f"GX-{n}"
    elif numbering == "roman_numerals":
        return _int_to_roman(n)
    elif numbering == "bates_numbered":
        return f"{n:03d}"
    elif numbering == "tab_numbers":
        return f"Tab {n}"
    elif numbering == "schedule_letters":
        return f"Schedule {_exhibit_letter(count)}"
    elif numbering == "appendix_letters":
        return f"Appendix {_exhibit_letter(count)}"
    else:  # "letters" default
        return _exhibit_letter(count)


def _exhibit_filename(exhibit_label: str, exhibit_name: str, ext: str, numbering: str = "letters") -> str:
    """Return the renamed filename for an exhibit — always prefixed with 'Exhibit'."""
    return f"Exhibit {exhibit_label} \u2014 {exhibit_name}{ext}"


def _next_exhibit_label(db, case_id: str, numbering: str = "letters") -> tuple:
    """Return (next_label, next_order) for a case's documents.
    Supports all 15 exhibit numbering formats."""
    row = db.execute(
        "SELECT COUNT(*) as cnt FROM documents WHERE case_id = ? AND exhibit_label IS NOT NULL",
        (case_id,)
    ).fetchone()
    count = row["cnt"] if row else 0
    return _make_exhibit_label(count, numbering), count


def _extract_text_from_pdf(file_bytes: bytes, max_chars: int = 500_000) -> str:
    """Extract text from PDF bytes using PyMuPDF — all pages, no truncation by default."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            page_text = page.get_text().strip()
            if page_text:
                pages.append(page_text)
        doc.close()
        text = "\n\n".join(pages)
        return text[:max_chars]
    except Exception as e:
        logger.warning(f"PDF text extraction failed: {e}")
        return ""


def _generate_html_from_docx(file_bytes: bytes) -> str:
    """Convert DOCX to HTML preserving fonts, sizes, bold/italic via python-docx."""
    try:
        import io as _io
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn
        import html as _html

        doc = Document(_io.BytesIO(file_bytes))
        parts = []

        def _run_style(run, para_style_font) -> str:
            styles = []
            # Font size
            sz = run.font.size or (para_style_font.size if para_style_font else None)
            if sz:
                styles.append(f"font-size:{sz.pt:.1f}pt")
            # Font family
            name = run.font.name or (para_style_font.name if para_style_font else None)
            if name:
                styles.append(f"font-family:'{name}',serif")
            # Bold
            b = run.bold if run.bold is not None else (para_style_font.bold if para_style_font else None)
            if b:
                styles.append("font-weight:bold")
            # Italic
            i = run.italic if run.italic is not None else (para_style_font.italic if para_style_font else None)
            if i:
                styles.append("font-style:italic")
            # Underline
            u = run.underline if run.underline is not None else (para_style_font.underline if para_style_font else None)
            if u:
                styles.append("text-decoration:underline")
            return ";".join(styles)

        for para in doc.paragraphs:
            raw = para.text
            sname = (para.style.name or "").lower()
            pf = para.paragraph_format
            psf = para.style.font if para.style else None

            # Choose tag
            if "heading 1" in sname:
                tag = "h1"
            elif "heading 2" in sname:
                tag = "h2"
            elif "heading 3" in sname:
                tag = "h3"
            elif "heading 4" in sname or "heading 5" in sname or "heading 6" in sname:
                tag = "h4"
            else:
                tag = "p"

            # Paragraph-level style
            p_styles = []
            if pf.left_indent:
                p_styles.append(f"margin-left:{pf.left_indent.pt:.0f}pt")
            if pf.right_indent:
                p_styles.append(f"margin-right:{pf.right_indent.pt:.0f}pt")
            if pf.space_before:
                p_styles.append(f"margin-top:{pf.space_before.pt:.0f}pt")
            if pf.space_after:
                p_styles.append(f"margin-bottom:{pf.space_after.pt:.0f}pt")
            if para.alignment is not None:
                align_map = {0: "left", 1: "center", 2: "right", 3: "justify"}
                al = align_map.get(para.alignment.value if hasattr(para.alignment, 'value') else para.alignment)
                if al:
                    p_styles.append(f"text-align:{al}")

            p_style_str = ";".join(p_styles)
            tag_open = f'<{tag} style="{p_style_str}">' if p_style_str else f'<{tag}>'

            if not raw.strip():
                parts.append(f"{tag_open}&nbsp;</{tag}>")
                continue

            # Build inline spans per run
            if para.runs:
                spans = []
                for run in para.runs:
                    if not run.text:
                        continue
                    rs = _run_style(run, psf)
                    text_esc = _html.escape(run.text)
                    if rs:
                        spans.append(f'<span style="{rs}">{text_esc}</span>')
                    else:
                        spans.append(text_esc)
                inner = "".join(spans)
            else:
                inner = _html.escape(raw)

            parts.append(f"{tag_open}{inner}</{tag}>")

        # Also process tables
        for table in doc.tables:
            rows_html = []
            for row in table.rows:
                cells_html = []
                for cell in row.cells:
                    cell_text = _html.escape(cell.text.strip())
                    cells_html.append(f'<td style="border:1px solid #d1d5db;padding:6px 10px;vertical-align:top">{cell_text}</td>')
                rows_html.append(f'<tr>{"".join(cells_html)}</tr>')
            parts.append(f'<table style="border-collapse:collapse;width:100%;margin:12px 0">{"".join(rows_html)}</table>')

        return "\n".join(parts)
    except Exception as e:
        logger.warning(f"DOCX HTML generation failed: {e}")
        return ""


def _generate_html_from_pdf(file_bytes: bytes) -> str:
    """Convert PDF to styled HTML preserving font sizes, bold, italic via PyMuPDF."""
    try:
        import fitz
        import html as _html

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts = []

        for page_num, page in enumerate(doc):
            if page_num > 0:
                parts.append('<hr style="border:none;border-top:2px dashed #d1d5db;margin:24px 0">')

            blocks = sorted(
                page.get_text("dict")["blocks"],
                key=lambda b: (round(b["bbox"][1] / 5) * 5, b["bbox"][0])
            )

            for block in blocks:
                if block.get("type") != 0:
                    continue

                lines_html = []
                block_max_size = 0

                for line in block.get("lines", []):
                    span_html = []
                    for span in line.get("spans", []):
                        text = span["text"]
                        if not text.strip():
                            continue
                        size = span["size"]
                        flags = span["flags"]
                        font = span["font"]
                        block_max_size = max(block_max_size, size)

                        styles = [f"font-size:{size:.1f}pt"]
                        if font and font != "(null)":
                            styles.append(f"font-family:'{font}',serif")
                        if flags & 16 or "Bold" in (font or ""):
                            styles.append("font-weight:bold")
                        if flags & 2 or "Italic" in (font or ""):
                            styles.append("font-style:italic")

                        span_html.append(
                            f'<span style="{";".join(styles)}">{_html.escape(text)}</span>'
                        )

                    if span_html:
                        lines_html.append("".join(span_html))

                if not lines_html:
                    continue

                inner = "<br>".join(lines_html)
                if block_max_size >= 16:
                    parts.append(f'<h2 style="font-size:{block_max_size:.1f}pt;margin:16px 0 8px">{inner}</h2>')
                elif block_max_size >= 13:
                    parts.append(f'<h3 style="font-size:{block_max_size:.1f}pt;margin:12px 0 6px">{inner}</h3>')
                else:
                    parts.append(f'<p style="font-size:{block_max_size:.1f}pt;margin:4px 0;line-height:1.6">{inner}</p>')

        doc.close()
        return "\n".join(parts)
    except Exception as e:
        logger.warning(f"PDF HTML generation failed: {e}")
        return ""


def _extract_text_from_docx(file_bytes: bytes, max_chars: int = 500_000) -> str:
    """Extract text from DOCX bytes, preserving paragraph boundaries with double newlines."""
    try:
        import zipfile
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        raw_paragraphs = tree.findall(".//w:p", ns)
        paras = []
        for p in raw_paragraphs:
            line = "".join(node.text or "" for node in p.findall(".//w:t", ns)).strip()
            if line:
                paras.append(line)
        text = "\n\n".join(paras)
        return text[:max_chars]
    except Exception as e:
        logger.warning(f"DOCX text extraction failed: {e}")
        return ""


def _ai_describe_image(file_bytes: bytes, original_filename: str, mime_type: str = "image/png") -> str:
    """Use GPT-5.4 vision to describe an image for exhibit naming."""
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return ""
    try:
        import base64
        from openai import OpenAI
        b64 = base64.b64encode(file_bytes).decode("utf-8")
        # Map common extensions to mime types
        ext = Path(original_filename).suffix.lower()
        mime_map = {
            ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/bmp",
            ".tiff": "image/tiff", ".tif": "image/tiff",
        }
        img_mime = mime_map.get(ext, mime_type)
        client = OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model=get_model_for_task("image_description"),
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a legal document analyst. Describe this image/screenshot in detail. "
                        "Focus on: any text visible, document type, parties mentioned, dates, "
                        "signatures, headers, or notable content. This will be used to generate "
                        "an exhibit name for a legal case. Be specific and factual."
                    ),
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": f"Describe this image (filename: {original_filename}). What does it show?"},
                        {"type": "image_url", "image_url": {"url": f"data:{img_mime};base64,{b64}", "detail": "high"}},
                    ],
                },
            ],
            max_completion_tokens=500,
            temperature=0.3,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logger.warning(f"AI image description failed: {e}")
        return ""


_EXHIBIT_NAME_SYSTEM_PROMPT = (
    "You are a legal document analyst. Given the text of a legal document, "
    "generate a short, descriptive exhibit name that MUST include:\n"
    "1. The document type (Agreement, Invoice, Email, Letter, Contract, etc.)\n"
    "2. The names of ALL signing parties or key parties mentioned\n"
    "3. The date the document was signed or executed (if found in the text)\n\n"
    "Format: '[Document Type] between [Party A] and [Party B] dated [Date] (signed by [Signatories])'\n\n"
    "Examples:\n"
    "- 'Consulting Services Agreement between ERC Ltd and Lebanon Medical dated May 17, 2023 (signed by Derek Cowan and Sarah Mills)'\n"
    "- 'Invoice #1234 from ABC Corp to XYZ Inc dated June 3, 2023'\n"
    "- 'Demand Letter from Smith & Associates to Jones Corp dated Jan 5, 2024'\n"
    "- 'Employment Contract between John Doe and Acme Corp dated March 1, 2022 (signed by John Doe and HR Director)'\n"
    "- 'Email correspondence between Smith and Jones re: payment terms (Jan 5-12, 2024)'\n\n"
    "CRITICAL RULES:\n"
    "- ALWAYS include party names if they appear anywhere in the document\n"
    "- ALWAYS include the signing/execution date if one appears in the document\n"
    "- If multiple dates exist, use the signing/execution date, not other dates\n"
    "- If no date is found, simply omit the date — do NOT write 'undated' or similar\n"
    "- If no party names are found, use whatever identifying info is available\n"
    "- Return ONLY the exhibit name, nothing else. Keep it under 150 characters."
)

_DOCUMENT_NAME_SYSTEM_PROMPT = (
    "You are a legal document analyst. Given the text of a document, "
    "generate a short, descriptive filename. Describe what the document is — "
    "include key parties and dates if present. "
    "This applies to ALL types of legal documents: court filings, motions, "
    "briefs, contracts, letters, orders, pleadings, discovery documents, "
    "general correspondence, and any other legal document. "
    "If no date is found, simply omit the date — do NOT write "
    "'date not found', 'undated', or any similar placeholder. "
    "Examples: 'Motion to Dismiss — Smith v. Jones', "
    "'Retainer Agreement with Johnson LLC dated March 2024', "
    "'Court Order Granting Summary Judgment', "
    "'Plaintiff Complaint — Jones v. City of Newark', "
    "'Subpoena Duces Tecum to ABC Corp', "
    "'Settlement Agreement — Johnson v. Metro Health'. "
    "Return ONLY the document name, nothing else. Keep it under 120 characters."
)


def _ai_generate_exhibit_name(text: str, original_filename: str) -> str:
    """Call OpenAI (with Anthropic fallback) to generate a descriptive exhibit name."""
    fallback_name = Path(original_filename).stem.replace("_", " ").replace("-", " ")
    if not text.strip():
        return fallback_name

    user_msg = f"Original filename: {original_filename}\n\nDocument text (first ~3000 chars):\n{text}"

    # Try OpenAI first
    openai_key = os.environ.get("OPENAI_API_KEY")
    if openai_key:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=openai_key)
            response = client.chat.completions.create(
                model=get_model_for_task("exhibit_naming"),
                messages=[
                    {"role": "system", "content": _EXHIBIT_NAME_SYSTEM_PROMPT},
                    {"role": "user", "content": user_msg},
                ],
                max_completion_tokens=200,
                temperature=0.3,
            )
            name = response.choices[0].message.content.strip().strip('"').strip("'")
            if name:
                return name
        except Exception as e:
            logger.warning(f"OpenAI exhibit naming failed: {e}")

    # Fallback to Anthropic (Claude)
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY")
    if anthropic_key:
        try:
            import anthropic
            client = anthropic.Anthropic(api_key=anthropic_key)
            response = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=200,
                system=_EXHIBIT_NAME_SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_msg}],
            )
            name = response.content[0].text.strip().strip('"').strip("'")
            if name:
                logger.info(f"Anthropic exhibit naming succeeded (OpenAI unavailable)")
                return name
        except Exception as e:
            logger.warning(f"Anthropic exhibit naming also failed: {e}")

    return fallback_name


def _ai_generate_document_name(text: str, original_filename: str) -> str:
    """Call OpenAI (with Anthropic fallback) to generate a clean descriptive name for a non-exhibit document."""
    fallback_name = Path(original_filename).stem.replace("_", " ").replace("-", " ")

    input_text = text.strip() if text else ""
    if not input_text:
        input_text = "(No text content extracted. Use the original filename to infer the document type.)"

    user_msg = f"Original filename: {original_filename}\n\nDocument text (first ~3000 chars):\n{input_text}"

    # Try OpenAI first
    openai_key = os.environ.get("OPENAI_API_KEY")
    if openai_key:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=openai_key)
            response = client.chat.completions.create(
                model=get_model_for_task("document_naming"),
                messages=[
                    {"role": "system", "content": _DOCUMENT_NAME_SYSTEM_PROMPT},
                    {"role": "user", "content": user_msg},
                ],
                max_completion_tokens=150,
                temperature=0.3,
            )
            name = response.choices[0].message.content.strip().strip('"').strip("'")
            if name:
                return name
        except Exception as e:
            logger.warning(f"OpenAI document naming failed: {e}")

    # Fallback to Anthropic (Claude)
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY")
    if anthropic_key:
        try:
            import anthropic
            client = anthropic.Anthropic(api_key=anthropic_key)
            response = client.messages.create(
                model="claude-sonnet-4-6",
                max_tokens=150,
                system=_DOCUMENT_NAME_SYSTEM_PROMPT,
                messages=[{"role": "user", "content": user_msg}],
            )
            name = response.content[0].text.strip().strip('"').strip("'")
            if name:
                logger.info(f"Anthropic document naming succeeded (OpenAI unavailable)")
                return name
        except Exception as e:
            logger.warning(f"Anthropic document naming also failed: {e}")

    return fallback_name


def _stamp_exhibit_on_pdf(file_bytes: bytes, exhibit_label: str, exhibit_name: str, numbering: str = "letters") -> bytes:
    """Insert a full exhibit cover page BEFORE the original document.
    Cover page layout:
      - 'EXHIBIT A' in very large bold red text (48pt)
      - '(Document Name)' in bold black text (28pt), in parentheses
    Original document pages follow unchanged."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")

        if len(doc) == 0:
            doc.close()
            return file_bytes

        # Use same dimensions as first page of original document
        first_page = doc[0]
        pw = first_page.rect.width
        ph = first_page.rect.height

        # Create new document: cover page + all original pages
        new_doc = fitz.open()

        # --- Cover Page ---
        cover = new_doc.new_page(width=pw, height=ph)

        # "EXHIBIT A" — large red, centered horizontally and ~35% from top
        exhibit_text = f"EXHIBIT {exhibit_label}"
        font = fitz.Font("helv")
        text_width = font.text_length(exhibit_text, fontsize=48)
        x_exhibit = (pw - text_width) / 2
        y_exhibit = ph * 0.38
        cover.insert_text(
            fitz.Point(x_exhibit, y_exhibit),
            exhibit_text,
            fontsize=48,
            fontname="helv",
            color=(1, 0, 0),  # RED
        )

        # "(Document Name)" — bold black, centered below exhibit label
        # Wrap long names into multiple lines
        if exhibit_name:
            name_fontsize = 28
            max_width = pw - 80  # 40px margin each side
            name_text = f"({exhibit_name})"

            # Simple word-wrap
            words = name_text.split()
            lines = []
            current_line = ""
            for word in words:
                test = f"{current_line} {word}".strip() if current_line else word
                if font.text_length(test, fontsize=name_fontsize) <= max_width:
                    current_line = test
                else:
                    if current_line:
                        lines.append(current_line)
                    current_line = word
            if current_line:
                lines.append(current_line)

            y_name = y_exhibit + 60  # gap below exhibit text
            for line in lines:
                line_width = font.text_length(line, fontsize=name_fontsize)
                x_name = (pw - line_width) / 2
                cover.insert_text(
                    fitz.Point(x_name, y_name),
                    line,
                    fontsize=name_fontsize,
                    fontname="helv",
                    color=(0, 0, 0),  # black
                )
                y_name += name_fontsize + 8  # line spacing

        # --- Copy all original pages unchanged ---
        # insert_pdf preserves rotation, landscape orientation, and all page properties
        new_doc.insert_pdf(doc)

        output = new_doc.tobytes()
        new_doc.close()
        doc.close()
        return output
    except Exception as e:
        logger.warning(f"PDF stamping failed: {e}")
        return file_bytes  # Return original if stamping fails


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff", ".tif", ".heic", ".heif"}


def _restamp_existing_exhibit(
    doc_id: str, case_id: str, tenant_id: str,
    file_bytes: bytes, filename: str, ext: str,
    file_path_on_disk: Path, exhibit_label: str, exhibit_name: str
):
    """Re-stamp a replaced exhibit file using its existing label and name.
    Never reassigns the label — preserves Exhibit A as A, F as F, etc."""
    try:
        # Extract fresh text/HTML from new file
        if ext == ".pdf":
            text = _extract_text_from_pdf(file_bytes)
            html = _generate_html_from_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            text = _extract_text_from_docx(file_bytes)
            html = _generate_html_from_docx(file_bytes)
        elif ext == ".txt":
            text = file_bytes.decode("utf-8", errors="ignore")[:500_000]
            import html as _html_mod
            html = "".join(f"<p>{_html_mod.escape(l)}</p>" for l in text.splitlines() if l.strip())
        else:
            text = ""
            html = ""

        numbering = "letters"
        with get_db() as db:
            row = db.execute("SELECT exhibit_numbering FROM cases WHERE id = ?", (case_id,)).fetchone()
            if row and row["exhibit_numbering"]:
                numbering = row["exhibit_numbering"]

        # Stamp PDF with EXISTING label and name — preserves landscape pages via insert_pdf
        if ext == ".pdf":
            stamped = _stamp_exhibit_on_pdf(file_bytes, exhibit_label, exhibit_name, numbering)
            file_path_on_disk.write_bytes(stamped)

        with get_db() as db:
            db.execute(
                "UPDATE documents SET content_text = ?, content_html = ? WHERE id = ?",
                (text, html, doc_id)
            )
        logger.info(f"Exhibit re-stamped: {doc_id} -> Exhibit {exhibit_label}")
    except Exception as e:
        logger.error(f"Exhibit re-stamp failed for {doc_id}: {e}\n{traceback.format_exc()}")


def _process_exhibit_async(doc_id: str, case_id: str, tenant_id: str,
                           file_bytes: bytes, original_filename: str, ext: str,
                           file_path_on_disk: Path, relative_path: str,
                           category: str = "evidence"):
    """Background task: extract text, AI name, stamp PDF, update DB.
    Only documents with category 'evidence' get exhibit labels and cover pages.
    All other categories just get an AI-generated descriptive filename."""
    try:
        # Look up case exhibit_numbering preference
        numbering = "letters"
        try:
            with get_db() as db:
                case_row = db.execute("SELECT exhibit_numbering FROM cases WHERE id = ?", (case_id,)).fetchone()
                if case_row and case_row["exhibit_numbering"]:
                    numbering = case_row["exhibit_numbering"]
        except Exception:
            pass  # column may not exist yet; default to letters

        # 1. Extract text and generate styled HTML
        if ext == ".pdf":
            text = _extract_text_from_pdf(file_bytes)
            html = _generate_html_from_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            text = _extract_text_from_docx(file_bytes)
            html = _generate_html_from_docx(file_bytes)
        elif ext == ".txt":
            text = file_bytes.decode("utf-8", errors="ignore")[:500_000]
            import html as _html_mod
            html = "".join(
                f"<p>{_html_mod.escape(line)}</p>"
                for line in text.splitlines() if line.strip()
            )
        elif ext in IMAGE_EXTENSIONS:
            text = _ai_describe_image(file_bytes, original_filename)
            html = f"<p>{text}</p>"
        else:
            text = ""
            html = ""

        is_evidence = (category == "evidence")

        if is_evidence:
            # --- EVIDENCE: full exhibit treatment (label, cover page, scheme-aware filename) ---
            with get_db() as db:
                exhibit_label, exhibit_order = _next_exhibit_label(db, case_id, numbering)
                exhibit_name = _ai_generate_exhibit_name(text, original_filename)
                new_filename = _exhibit_filename(exhibit_label, exhibit_name, ext, numbering)

                # Stamp PDF with cover page
                if ext == ".pdf":
                    stamped_bytes = _stamp_exhibit_on_pdf(file_bytes, exhibit_label, exhibit_name, numbering)
                    file_path_on_disk.write_bytes(stamped_bytes)

                db.execute(
                    "UPDATE documents SET exhibit_label = ?, exhibit_name = ?, exhibit_order = ?, filename = ?, content_text = ?, content_html = ? WHERE id = ?",
                    (exhibit_label, exhibit_name, exhibit_order, new_filename, text, html, doc_id)
                )
                logger.info(f"Exhibit processed: {doc_id} -> {new_filename}")
        else:
            # --- NON-EVIDENCE: just AI-rename, no exhibit label or stamp ---
            doc_name = _ai_generate_document_name(text, original_filename)
            new_filename = f"{doc_name}{ext}"
            with get_db() as db:
                db.execute(
                    "UPDATE documents SET filename = ?, content_text = ?, content_html = ? WHERE id = ?",
                    (new_filename, text, html, doc_id)
                )
                logger.info(f"Document renamed: {doc_id} -> {new_filename}")

    except Exception as e:
        logger.error(f"Document processing failed for {doc_id}: {e}\n{traceback.format_exc()}")
        # Fallback for evidence: still try to assign a label even if AI failed
        if category == "evidence":
            try:
                with get_db() as db:
                    exhibit_label, exhibit_order = _next_exhibit_label(db, case_id, numbering)
                    fallback_name = Path(original_filename).stem.replace("_", " ").replace("-", " ")
                    new_filename = _exhibit_filename(exhibit_label, fallback_name, ext, numbering)
                    db.execute(
                        "UPDATE documents SET exhibit_label = ?, exhibit_name = ?, exhibit_order = ?, filename = ? WHERE id = ?",
                        (exhibit_label, fallback_name, exhibit_order, new_filename, doc_id)
                    )
            except Exception:
                pass


# Pydantic models for exhibit endpoints
class ExhibitUpdate(BaseModel):
    exhibit_name: Optional[str] = None
    exhibit_label: Optional[str] = None

class DocumentRename(BaseModel):
    filename: str

class DocumentCategoryUpdate(BaseModel):
    category: str

class ExhibitReorderItem(BaseModel):
    doc_id: str
    order: int

class ExhibitReorderRequest(BaseModel):
    items: List[ExhibitReorderItem]

class DocumentMergeRequest(BaseModel):
    doc_ids: List[str] = []  # empty = merge all
    merge_all: bool = False
    filename: str = "Merged Document"
    category: str = "evidence"
    bates_prefix: str = ""  # e.g. "BATES" — empty = no Bates numbering
    bates_start: int = 1

router = APIRouter(prefix="/api/cases", tags=["cases"])


def _calculate_urgency_score(case: dict) -> float:
    """Calculate urgency score based on deadline proximity vs completion percentage."""
    score = 50.0  # base score

    if case.get("filing_deadline"):
        try:
            deadline = datetime.fromisoformat(case["filing_deadline"])
            now = datetime.now(timezone.utc)
            days_remaining = (deadline - now).days
            if days_remaining < 0:
                score += 50  # Overdue
            elif days_remaining < 3:
                score += 40
            elif days_remaining < 7:
                score += 30
            elif days_remaining < 14:
                score += 20
            elif days_remaining < 30:
                score += 10
        except (ValueError, TypeError):
            pass

    completion = case.get("completion_percentage", 0) or 0
    score -= completion * 0.3  # Higher completion reduces urgency

    priority_weights = {"critical": 20, "high": 15, "medium": 5, "low": 0}
    score += priority_weights.get(case.get("priority", "medium"), 5)

    return max(0, min(100, score))


@router.get("")
async def list_cases(
    status: str = None,
    case_type: str = None,
    sort_by: str = "urgency_score",
    current_user: dict = Depends(get_current_user)
):
    """List cases filtered by tenant."""
    tenant_id = current_user["tenant_id"]
    with get_db() as db:
        query = "SELECT * FROM cases WHERE tenant_id = ?"
        params: list = [tenant_id]
        if status:
            query += " AND status = ?"
            params.append(status)
        if case_type:
            query += " AND case_type = ?"
            params.append(case_type)

        if sort_by == "urgency_score":
            query += " ORDER BY urgency_score DESC"
        elif sort_by == "filing_deadline":
            query += " ORDER BY filing_deadline ASC"
        elif sort_by == "created_at":
            query += " ORDER BY created_at DESC"
        else:
            query += " ORDER BY urgency_score DESC"

        cases = db.execute(query, params).fetchall()
        result = []
        for c in cases:
            case_dict = dict(c)
            # Recalculate urgency score
            new_score = _calculate_urgency_score(case_dict)
            if abs(new_score - (case_dict.get("urgency_score") or 0)) > 1:
                db.execute(
                    "UPDATE cases SET urgency_score = ? WHERE id = ?",
                    (new_score, case_dict["id"])
                )
                case_dict["urgency_score"] = new_score

            # Get task stats
            tasks = db.execute(
                "SELECT status, COUNT(*) as cnt FROM tasks WHERE case_id = ? GROUP BY status",
                (case_dict["id"],)
            ).fetchall()
            task_stats = {t["status"]: t["cnt"] for t in tasks}
            case_dict["task_stats"] = task_stats
            total = sum(task_stats.values())
            completed = task_stats.get("completed", 0)
            if total > 0:
                case_dict["completion_percentage"] = round((completed / total) * 100, 1)

            result.append(case_dict)
        return result


@router.post("")
async def create_case(req: CaseCreate, current_user: dict = Depends(get_current_user)):
    """Create a new case."""
    tenant_id = current_user["tenant_id"]
    case_id = generate_id()
    now = datetime.now(timezone.utc).isoformat()

    case_number = (req.case_number or "").strip() or None

    exhibit_numbering = req.exhibit_numbering if req.exhibit_numbering in EXHIBIT_SCHEMES else "letters"

    case_data = {
        "id": case_id,
        "tenant_id": tenant_id,
        "title": req.title,
        "case_number": case_number,
        "case_type": req.case_type,
        "description": req.description,
        "client_name": req.client_name,
        "opposing_party": req.opposing_party,
        "court": req.court,
        "judge": req.judge,
        "filing_deadline": req.filing_deadline,
        "trial_date": req.trial_date,
        "uscis_receipt_number": req.uscis_receipt_number,
        "priority": req.priority,
        "exhibit_numbering": exhibit_numbering,
        "created_by": current_user["sub"],
        "assigned_attorney_id": current_user["sub"],
        "created_at": now,
        "updated_at": now,
    }

    urgency = _calculate_urgency_score(case_data)
    case_data["urgency_score"] = urgency

    with get_db() as db:
        db.execute(
            """INSERT INTO cases (id, tenant_id, title, case_number, case_type, description,
               client_name, opposing_party, court, judge, filing_deadline, trial_date,
               uscis_receipt_number, priority, urgency_score, exhibit_numbering, created_by, assigned_attorney_id,
               created_at, updated_at)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (case_id, tenant_id, req.title, case_number, req.case_type, req.description,
             req.client_name, req.opposing_party, req.court, req.judge, req.filing_deadline,
             req.trial_date, req.uscis_receipt_number, req.priority, urgency, exhibit_numbering,
             current_user["sub"], current_user["sub"], now, now)
        )

        # Auto-populate tasks from workflow template if case_type matches
        template = db.execute(
            "SELECT * FROM workflow_templates WHERE case_type = ?",
            (req.case_type,)
        ).fetchone()
        if template:
            tasks = json.loads(template["tasks_json"])
            for i, task_title in enumerate(tasks):
                task_id = generate_id()
                db.execute(
                    """INSERT INTO tasks (id, case_id, tenant_id, title, status, priority)
                       VALUES (?, ?, ?, ?, 'pending', 'medium')""",
                    (task_id, case_id, tenant_id, task_title)
                )

    case_data["task_stats"] = {}
    return case_data


@router.get("/{case_id}")
async def get_case(case_id: str, current_user: dict = Depends(get_current_user)):
    """Get case details with tasks, documents, and timeline."""
    tenant_id = current_user["tenant_id"]
    with get_db() as db:
        case = db.execute(
            "SELECT * FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, tenant_id)
        ).fetchone()
        if not case:
            # Check if user has expert access
            access = db.execute(
                """SELECT ce.* FROM case_experts ce
                   JOIN cases c ON ce.case_id = c.id
                   WHERE ce.case_id = ? AND ce.expert_id = ? AND ce.status = 'active'""",
                (case_id, current_user["sub"])
            ).fetchone()
            if not access:
                raise HTTPException(status_code=404, detail="Case not found")
            case = db.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()

        case_dict = dict(case)

        # Get tasks
        tasks = db.execute(
            "SELECT * FROM tasks WHERE case_id = ? ORDER BY created_at ASC",
            (case_id,)
        ).fetchall()
        case_dict["tasks"] = [dict(t) for t in tasks]

        # Get documents (ordered by exhibit_order, then created_at)
        docs = db.execute(
            "SELECT * FROM documents WHERE case_id = ? ORDER BY exhibit_order ASC, created_at ASC",
            (case_id,)
        ).fetchall()
        case_dict["documents"] = [dict(d) for d in docs]

        # Get timeline events
        timeline = db.execute(
            "SELECT * FROM case_timeline WHERE case_id = ? ORDER BY event_date ASC",
            (case_id,)
        ).fetchall()
        case_dict["timeline"] = [dict(t) for t in timeline]

        # Get contradictions
        contradictions = db.execute(
            "SELECT * FROM contradictions WHERE case_id = ? ORDER BY severity DESC",
            (case_id,)
        ).fetchall()
        case_dict["contradictions"] = [dict(c) for c in contradictions]

        # Get assigned experts
        experts = db.execute(
            """SELECT ce.*, u.full_name, u.email, u.specializations
               FROM case_experts ce JOIN users u ON ce.expert_id = u.id
               WHERE ce.case_id = ? AND ce.status = 'active'""",
            (case_id,)
        ).fetchall()
        case_dict["experts"] = [dict(e) for e in experts]

        return case_dict


@router.patch("/{case_id}")
async def update_case(
    case_id: str,
    req: CaseUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Update case details."""
    tenant_id = current_user["tenant_id"]
    with get_db() as db:
        case = db.execute(
            "SELECT * FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, tenant_id)
        ).fetchone()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")

        updates = {}
        nullable_fields = {
            "case_number",
            "description",
            "client_name",
            "opposing_party",
            "court",
            "judge",
            "filing_deadline",
            "trial_date",
            "uscis_receipt_number",
            "assigned_attorney_id",
            "exhibit_numbering",
        }
        for field, value in req.model_dump(exclude_unset=True).items():
            # Allow explicitly clearing nullable fields by sending null
            if value is not None or field in nullable_fields:
                updates[field] = value

        if not updates:
            return dict(case)

        updates["updated_at"] = datetime.now(timezone.utc).isoformat()

        set_clause = ", ".join(f"{k} = ?" for k in updates.keys())
        values = list(updates.values()) + [case_id, tenant_id]
        db.execute(
            f"UPDATE cases SET {set_clause} WHERE id = ? AND tenant_id = ?",
            values
        )

        # Recalculate urgency
        updated = db.execute("SELECT * FROM cases WHERE id = ?", (case_id,)).fetchone()
        updated_dict = dict(updated)
        new_score = _calculate_urgency_score(updated_dict)
        db.execute("UPDATE cases SET urgency_score = ? WHERE id = ?", (new_score, case_id))
        updated_dict["urgency_score"] = new_score

        return updated_dict


@router.delete("/{case_id}")
async def delete_case(case_id: str, current_user: dict = Depends(get_current_user)):
    """Delete a case and all associated data."""
    tenant_id = current_user["tenant_id"]
    with get_db() as db:
        case = db.execute(
            "SELECT * FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, tenant_id)
        ).fetchone()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")

        # Delete uploaded files from disk (best effort)
        docs = db.execute(
            "SELECT file_path FROM documents WHERE case_id = ?",
            (case_id,)
        ).fetchall()
        for d in docs:
            try:
                fp = d["file_path"]
                if fp and os.path.exists(fp):
                    os.remove(fp)
            except Exception:
                # Don't block deletion if file cleanup fails
                pass

        # Delete drafts + dependent records
        draft_rows = db.execute(
            "SELECT id FROM legal_drafts WHERE case_id = ?",
            (case_id,)
        ).fetchall()
        draft_ids = [r["id"] for r in draft_rows]
        for did in draft_ids:
            db.execute("DELETE FROM draft_versions WHERE draft_id = ?", (did,))
            db.execute("DELETE FROM research_citations WHERE draft_id = ?", (did,))
            db.execute("DELETE FROM draft_comments WHERE draft_id = ?", (did,))
        db.execute("DELETE FROM legal_drafts WHERE case_id = ?", (case_id,))

        # Delete all associated data
        db.execute("DELETE FROM tasks WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM time_entries WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM documents WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM case_timeline WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM contradictions WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM case_experts WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM waitlist WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM discovery_items WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM witnesses WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM chat_messages WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM case_contacts WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM case_emails WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM case_notes WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM case_pipeline WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM case_access WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM timeline_events WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM contracts WHERE case_id = ?", (case_id,))
        db.execute("DELETE FROM cases WHERE id = ? AND tenant_id = ?", (case_id, tenant_id))

    return {"ok": True, "message": "Case deleted"}


# Task sub-routes
@router.get("/{case_id}/tasks")
async def get_tasks(case_id: str, current_user: dict = Depends(get_current_user)):
    """Get all tasks for a case."""
    with get_db() as db:
        tasks = db.execute(
            "SELECT * FROM tasks WHERE case_id = ? ORDER BY created_at ASC",
            (case_id,)
        ).fetchall()
        return [dict(t) for t in tasks]


@router.post("/{case_id}/tasks")
async def create_task(
    case_id: str,
    req: TaskCreate,
    current_user: dict = Depends(get_current_user)
):
    """Create a task for a case."""
    tenant_id = current_user["tenant_id"]
    task_id = generate_id()

    with get_db() as db:
        # Verify case exists
        case = db.execute(
            "SELECT id FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, tenant_id)
        ).fetchone()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")

        db.execute(
            """INSERT INTO tasks (id, case_id, tenant_id, title, description, assigned_to,
               due_date, priority, parent_task_id)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (task_id, case_id, tenant_id, req.title, req.description, req.assigned_to,
             req.due_date, req.priority, req.parent_task_id)
        )
        return {
            "id": task_id, "case_id": case_id, "title": req.title,
            "status": "pending", "priority": req.priority
        }


@router.patch("/tasks/{task_id}")
async def update_task(
    task_id: str,
    req: TaskUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Update task status/details."""
    with get_db() as db:
        task = db.execute(
            "SELECT * FROM tasks WHERE id = ? AND tenant_id = ?",
            (task_id, current_user["tenant_id"])
        ).fetchone()
        if not task:
            raise HTTPException(status_code=404, detail="Task not found")

        updates = {}
        for field, value in req.model_dump(exclude_unset=True).items():
            if value is not None:
                updates[field] = value

        if "status" in updates and updates["status"] == "completed":
            updates["completed_at"] = datetime.now(timezone.utc).isoformat()

        if not updates:
            return dict(task)

        set_clause = ", ".join(f"{k} = ?" for k in updates.keys())
        values = list(updates.values()) + [task_id]
        db.execute(f"UPDATE tasks SET {set_clause} WHERE id = ?", values)

        updated = db.execute("SELECT * FROM tasks WHERE id = ?", (task_id,)).fetchone()
        return dict(updated)


# Document sub-routes
@router.get("/{case_id}/documents")
async def get_documents(case_id: str, current_user: dict = Depends(get_current_user)):
    """List documents for a case."""
    with get_db() as db:
        docs = db.execute(
            "SELECT * FROM documents WHERE case_id = ? ORDER BY created_at DESC",
            (case_id,)
        ).fetchall()
        return [dict(d) for d in docs]


@router.post("/{case_id}/documents")
async def create_document(
    case_id: str,
    req: DocumentCreate,
    current_user: dict = Depends(get_current_user)
):
    """Create a document record (file upload handled separately via pre-signed URL)."""
    tenant_id = current_user["tenant_id"]
    doc_id = generate_id()
    file_path = f"tenants/{tenant_id}/cases/{case_id}/{doc_id}/{req.filename}"

    with get_db() as db:
        db.execute(
            """INSERT INTO documents (id, case_id, tenant_id, filename, file_path, category,
               uploaded_by, content_text)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (doc_id, case_id, tenant_id, req.filename, file_path, req.category,
             current_user["sub"], req.content_text)
        )
        return {
            "id": doc_id,
            "case_id": case_id,
            "filename": req.filename,
            "file_path": file_path,
            "category": req.category,
        }


# File upload endpoint
@router.post("/{case_id}/documents/upload")
async def upload_document_file(
    case_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    category: str = Form(default="general"),
    notes: str = Form(default=""),
    current_user: dict = Depends(get_current_user)
):
    """Upload an actual file to a case. Auto-assigns exhibit label and AI-generated name."""
    tenant_id = current_user["tenant_id"]

    # Validate file extension
    filename = file.filename or "unnamed_file"
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"File type '{ext}' not allowed. Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}")

    # Read file content
    file_bytes = await file.read()
    file_size = len(file_bytes)
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB")
    if file_size == 0:
        raise HTTPException(status_code=400, detail="File is empty")

    doc_id = generate_id()
    mime_type = file.content_type or "application/octet-stream"

    # Create directory and save file
    upload_dir = Path(UPLOAD_BASE_DIR) / tenant_id / case_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_filename = f"{doc_id}_{filename}"
    file_path = upload_dir / safe_filename
    file_path.write_bytes(file_bytes)

    # Store relative path in DB
    relative_path = f"{tenant_id}/{case_id}/{safe_filename}"

    with get_db() as db:
        # Verify case exists
        case = db.execute(
            "SELECT id FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, tenant_id)
        ).fetchone()
        if not case:
            # Clean up file
            file_path.unlink(missing_ok=True)
            raise HTTPException(status_code=404, detail="Case not found")

        db.execute(
            """INSERT INTO documents (id, case_id, tenant_id, filename, file_path, file_size, mime_type,
               category, uploaded_by, content_text)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (doc_id, case_id, tenant_id, filename, relative_path, file_size, mime_type,
             category, current_user["sub"], notes)
        )

    # Process document in background (AI naming; exhibit labeling only for evidence)
    background_tasks.add_task(
        _process_exhibit_async, doc_id, case_id, tenant_id,
        file_bytes, filename, ext, file_path, relative_path,
        category
    )

    # Deduct 1 credit for the AI naming/tagging task (nano model, negligible cost)
    with get_db() as _db:
        _up_row = _db.execute("SELECT subscription_status FROM users WHERE id=?", (current_user["sub"],)).fetchone()
        if _up_row:
            deduct_credits(current_user["sub"], _up_row["subscription_status"], 1, "document_naming", _db)

    return {
        "id": doc_id,
        "case_id": case_id,
        "filename": filename,
        "file_size": file_size,
        "mime_type": mime_type,
        "category": category,
        "has_file": True,
        "exhibit_processing": category == "evidence",
    }


# File download endpoint
@router.get("/debug/upload-config")
async def debug_upload_config(current_user: dict = Depends(get_current_user)):
    """Diagnostic: show upload directory config and what files exist on disk."""
    import os
    base = Path(UPLOAD_BASE_DIR)
    entries = []
    if base.exists():
        try:
            for root, dirs, files in os.walk(str(base)):
                for f in files:
                    full = Path(root) / f
                    rel = full.relative_to(base)
                    entries.append(str(rel))
                if len(entries) > 200:
                    entries.append("... (truncated at 200 files)")
                    break
        except Exception as e:
            entries = [f"walk error: {e}"]
    return {
        "UPLOAD_BASE_DIR": str(base),
        "base_exists": base.exists(),
        "base_is_dir": base.is_dir() if base.exists() else False,
        "files_on_disk": entries,
    }


@router.get("/documents/{doc_id}/debug-path")
async def debug_document_path(doc_id: str, current_user: dict = Depends(get_current_user)):
    """Diagnostic: show exactly what path would be used to serve this document."""
    with get_db() as db:
        doc = db.execute(
            "SELECT id, filename, file_path, mime_type FROM documents WHERE id = ? AND tenant_id = ?",
            (doc_id, current_user["tenant_id"])
        ).fetchone()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        stored_path = doc["file_path"] or ""
        resolved = str(Path(UPLOAD_BASE_DIR) / stored_path) if stored_path else None
        exists = Path(resolved).exists() if resolved else False
        return {
            "doc_id": doc_id,
            "filename": doc["filename"],
            "stored_path_in_db": stored_path,
            "UPLOAD_BASE_DIR": UPLOAD_BASE_DIR,
            "resolved_full_path": resolved,
            "file_exists": exists,
        }


@router.get("/documents/{doc_id}/download")
async def download_document_file(
    doc_id: str,
    current_user: dict = Depends(get_download_user)
):
    """Download a document file."""
    with get_db() as db:
        doc = db.execute(
            "SELECT * FROM documents WHERE id = ? AND tenant_id = ?",
            (doc_id, current_user["tenant_id"])
        ).fetchone()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        stored_path = doc["file_path"] or ""
        file_path = Path(UPLOAD_BASE_DIR) / stored_path if stored_path else None

        # Fallback: if stored path is already absolute and exists, use it directly
        if (file_path is None or not file_path.exists()) and stored_path.startswith("/"):
            candidate = Path(stored_path)
            if candidate.exists():
                file_path = candidate

        if file_path is None or not file_path.exists():
            logger.warning(
                "Download 404: doc=%s UPLOAD_BASE_DIR=%r stored_path=%r resolved=%r",
                doc_id, UPLOAD_BASE_DIR, stored_path, str(file_path)
            )
            raise HTTPException(
                status_code=404,
                detail=f"File not found on server. base={UPLOAD_BASE_DIR!r} stored={stored_path!r}"
            )

        return FileResponse(
            path=str(file_path),
            filename=doc["filename"],
            media_type=doc["mime_type"] or "application/octet-stream",
        )


@router.get("/{case_id}/documents/download-all")
async def download_all_documents_as_pdf(
    case_id: str,
    bates_prefix: str = "",
    bates_start: int = 1,
    current_user: dict = Depends(get_download_user)
):
    """Merge ALL case documents (every file type) into one PDF for download.
    Optional Bates numbering via query params: ?bates_prefix=BATES&bates_start=1"""
    import fitz
    from fastapi.responses import Response

    tenant_id = current_user["tenant_id"]
    with get_db() as db:
        case = db.execute(
            "SELECT * FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, tenant_id)
        ).fetchone()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")

        docs = db.execute(
            "SELECT * FROM documents WHERE case_id = ? AND tenant_id = ? ORDER BY exhibit_order ASC, created_at ASC",
            (case_id, tenant_id)
        ).fetchall()
        if not docs:
            raise HTTPException(status_code=404, detail="No documents in this case")
        docs = [dict(d) for d in docs]

    # Sort: court_filing FIRST (main document), then petition, then exhibits by label (A, B, C...)
    # This ensures Statement of Claim / main filing is always on top
    def _doc_sort_key(d):
        label = d.get('exhibit_label') or ''
        cat = d.get('category', '')
        if cat == 'court_filing':
            return (0, '', d.get('created_at', ''))
        elif cat == 'petition':
            return (1, '', d.get('created_at', ''))
        elif cat == 'evidence' and label:
            # Evidence with exhibit label: after main filings, by label alphabetically
            return (2, label, d.get('created_at', ''))
        elif cat == 'evidence':
            # Evidence without label yet: after labeled evidence
            return (2, 'ZZZZ', d.get('created_at', ''))
        else:
            return (3, '', d.get('created_at', ''))
    docs.sort(key=_doc_sort_key)

    merged_pdf = fitz.open()
    for doc in docs:
        if not doc.get("file_path"):
            continue
        fp = Path(UPLOAD_BASE_DIR) / doc["file_path"]
        if not fp.exists():
            continue
        _add_file_to_pdf(merged_pdf, fp, doc.get("filename", ""))

    if merged_pdf.page_count == 0:
        merged_pdf.close()
        raise HTTPException(status_code=400, detail="No mergeable content found")

    # Apply Bates numbering if requested
    if bates_prefix:
        _apply_bates_numbering(merged_pdf, bates_prefix, bates_start)

    merged_bytes = merged_pdf.tobytes()
    merged_pdf.close()

    case_title = dict(case).get("title", "Case").replace(" ", "_")
    filename = f"{case_title}_Complete_Filing.pdf"

    return Response(
        content=merged_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@router.get("/{case_id}/documents/download-zip")
async def download_all_documents_as_zip(
    case_id: str,
    current_user: dict = Depends(get_download_user)
):
    """Download ALL case documents as a ZIP archive, preserving original file formats."""
    import zipfile
    from fastapi.responses import Response

    tenant_id = current_user["tenant_id"]
    with get_db() as db:
        case = db.execute(
            "SELECT * FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, tenant_id)
        ).fetchone()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")

        docs = db.execute(
            "SELECT * FROM documents WHERE case_id = ? AND tenant_id = ? ORDER BY exhibit_order ASC, created_at ASC",
            (case_id, tenant_id)
        ).fetchall()
        if not docs:
            raise HTTPException(status_code=404, detail="No documents in this case")
        docs = [dict(d) for d in docs]

    # Sort: court_filing FIRST (main document), then petition, then exhibits by label (A, B, C...)
    def _zip_sort_key(d):
        label = d.get('exhibit_label') or ''
        cat = d.get('category', '')
        if cat == 'court_filing':
            return (0, '', d.get('created_at', ''))
        elif cat == 'petition':
            return (1, '', d.get('created_at', ''))
        elif cat == 'evidence' and label:
            return (2, label, d.get('created_at', ''))
        elif cat == 'evidence':
            return (2, 'ZZZZ', d.get('created_at', ''))
        else:
            return (3, '', d.get('created_at', ''))
    docs.sort(key=_zip_sort_key)

    buf = io.BytesIO()
    used_names: dict = {}
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for doc in docs:
            if not doc.get("file_path"):
                continue
            fp = Path(UPLOAD_BASE_DIR) / doc["file_path"]
            if not fp.exists():
                continue
            # Use display filename, deduplicate if needed
            name = doc.get("filename") or fp.name
            if name in used_names:
                used_names[name] += 1
                stem = Path(name).stem
                ext = Path(name).suffix
                name = f"{stem} ({used_names[name]}){ext}"
            else:
                used_names[name] = 1
            zf.write(str(fp), name)

    zip_bytes = buf.getvalue()
    case_title = dict(case).get("title", "Case").replace(" ", "_")
    filename = f"{case_title}_Documents.zip"

    return Response(
        content=zip_bytes,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


# Delete document endpoint
@router.delete("/documents/{doc_id}")
async def delete_document(
    doc_id: str,
    current_user: dict = Depends(get_current_user)
):
    """Delete a document and its file."""
    with get_db() as db:
        doc = db.execute(
            "SELECT * FROM documents WHERE id = ? AND tenant_id = ?",
            (doc_id, current_user["tenant_id"])
        ).fetchone()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        # Delete file from disk
        if doc["file_path"]:
            file_path = Path(UPLOAD_BASE_DIR) / doc["file_path"]
            file_path.unlink(missing_ok=True)

        db.execute("DELETE FROM documents WHERE id = ?", (doc_id,))

    return {"status": "deleted", "document_id": doc_id}


@router.patch("/documents/{doc_id}/rename")
async def rename_document(
    doc_id: str,
    req: DocumentRename,
    current_user: dict = Depends(get_current_user)
):
    """Rename a document's display filename."""
    new_name = (req.filename or "").strip()
    if not new_name:
        raise HTTPException(status_code=400, detail="Filename cannot be empty")

    with get_db() as db:
        doc = db.execute(
            "SELECT * FROM documents WHERE id = ? AND tenant_id = ?",
            (doc_id, current_user["tenant_id"])
        ).fetchone()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        db.execute("UPDATE documents SET filename = ? WHERE id = ?", (new_name, doc_id))
        updated = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
        return dict(updated)


class DocumentAIRenameRequest(BaseModel):
    instructions: Optional[str] = ""


def _ai_rename_with_instructions(text: str, original_filename: str, instructions: str, is_exhibit: bool) -> str:
    """AI rename driven by user-supplied instructions."""
    fallback = Path(original_filename).stem.replace("_", " ").replace("-", " ")
    input_text = text.strip() if text else "(No text content extracted.)"
    user_msg = (
        f"Original filename: {original_filename}\n\n"
        f"User instruction: {instructions}\n\n"
        f"Document text (first ~2000 chars):\n{input_text[:2000]}"
    )
    system_msg = (
        "You are a legal document labeler. Follow the user instruction as the primary directive. "
        "Rules: 1. Follow the user instruction — it overrides everything else. "
        "2. Keep the name short (under 80 characters). "
        "3. Do NOT include a date unless the instruction or document specifically mentions one. "
        "4. Do NOT write placeholder text like 'undated' or 'not stated'. "
        "Return ONLY the name, nothing else."
    )
    openai_key = os.environ.get("OPENAI_API_KEY")
    if openai_key:
        try:
            from openai import OpenAI
            client = OpenAI(api_key=openai_key)
            resp = client.chat.completions.create(
                model=get_model_for_task("exhibit_naming" if is_exhibit else "document_naming"),
                messages=[{"role": "system", "content": system_msg}, {"role": "user", "content": user_msg}],
                max_completion_tokens=150, temperature=0.3,
            )
            name = resp.choices[0].message.content.strip().strip('"').strip("'")
            if name:
                return name
        except Exception as e:
            logger.warning(f"AI rename instructions (OpenAI) failed: {e}")
    anthropic_key = os.environ.get("ANTHROPIC_API_KEY")
    if anthropic_key:
        try:
            import anthropic as _anthropic
            client = _anthropic.Anthropic(api_key=anthropic_key)
            resp = client.messages.create(
                model="claude-sonnet-4-6", max_tokens=150,
                system=system_msg, messages=[{"role": "user", "content": user_msg}],
            )
            name = resp.content[0].text.strip().strip('"').strip("'")
            if name:
                return name
        except Exception as e:
            logger.warning(f"AI rename instructions (Anthropic) failed: {e}")
    return fallback


def _restamp_exhibit_pdf(file_bytes: bytes, exhibit_label: str, new_name: str, numbering: str) -> bytes:
    """Remove the existing cover page then re-stamp with a new name."""
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        if len(doc) < 2:
            doc.close()
            return _stamp_exhibit_on_pdf(file_bytes, exhibit_label, new_name, numbering)
        doc.delete_page(0)
        stripped = doc.tobytes()
        doc.close()
        return _stamp_exhibit_on_pdf(stripped, exhibit_label, new_name, numbering)
    except Exception as e:
        logger.warning(f"Cover page re-stamp failed: {e}")
        return file_bytes


@router.post("/documents/{doc_id}/ai-rename")
async def ai_rename_document(
    doc_id: str,
    req: DocumentAIRenameRequest,
    current_user: dict = Depends(get_current_user)
):
    """AI-rename a document, optionally guided by user instructions.
    For exhibits: updates exhibit_name + filename and re-stamps PDF cover page.
    For ordinary docs: updates filename only."""
    with get_db() as db:
        doc = db.execute(
            "SELECT * FROM documents WHERE id = ? AND tenant_id = ?",
            (doc_id, current_user["tenant_id"])
        ).fetchone()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        doc_dict = dict(doc)

    text = doc_dict.get("content_text") or ""
    original_filename = doc_dict.get("filename") or ""
    exhibit_label = doc_dict.get("exhibit_label") or ""
    file_path_str = doc_dict.get("file_path") or ""
    ext = Path(file_path_str).suffix.lower() if file_path_str else ""
    instructions = (req.instructions or "").strip()

    if exhibit_label:
        new_name = (
            _ai_rename_with_instructions(text, original_filename, instructions, is_exhibit=True)
            if instructions else _ai_generate_exhibit_name(text, original_filename)
        )
        numbering = "letters"
        with get_db() as db:
            row = db.execute("SELECT exhibit_numbering FROM cases WHERE id = ?", (doc_dict["case_id"],)).fetchone()
            if row and row["exhibit_numbering"]:
                numbering = row["exhibit_numbering"]

        new_filename = _exhibit_filename(exhibit_label, new_name, ext, numbering)
        fp = Path(UPLOAD_BASE_DIR) / file_path_str
        if ext == ".pdf" and fp.exists():
            try:
                stamped = _restamp_exhibit_pdf(fp.read_bytes(), exhibit_label, new_name, numbering)
                fp.write_bytes(stamped)
            except Exception as e:
                logger.warning(f"AI rename re-stamp failed for {doc_id}: {e}")

        with get_db() as db:
            db.execute("UPDATE documents SET exhibit_name = ?, filename = ? WHERE id = ?",
                       (new_name, new_filename, doc_id))
            updated = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
        return dict(updated)
    else:
        new_name = (
            _ai_rename_with_instructions(text, original_filename, instructions, is_exhibit=False)
            if instructions else _ai_generate_document_name(text, original_filename)
        )
        with get_db() as db:
            db.execute("UPDATE documents SET filename = ? WHERE id = ?", (new_name, doc_id))
            updated = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
        return dict(updated)


@router.post("/documents/{doc_id}/replace-file")
async def replace_document_file(
    doc_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user: dict = Depends(get_current_user)
):
    """Replace the physical file for an existing document record.
    Re-runs exhibit AI processing if the document is evidence."""
    tenant_id = current_user["tenant_id"]

    with get_db() as db:
        doc = db.execute(
            "SELECT * FROM documents WHERE id = ? AND tenant_id = ?",
            (doc_id, tenant_id)
        ).fetchone()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        doc_dict = dict(doc)

    # For exhibits that already have a nice AI-generated filename, preserve it
    upload_filename = file.filename or doc_dict["filename"]
    ext = Path(upload_filename).suffix.lower()
    filename = doc_dict["filename"] if doc_dict.get("exhibit_label") else upload_filename
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"File type '{ext}' not allowed")

    file_bytes = await file.read()
    file_size = len(file_bytes)
    if file_size == 0:
        raise HTTPException(status_code=400, detail="File is empty")
    if file_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large (max 100MB)")

    mime_type = file.content_type or "application/octet-stream"
    case_id = doc_dict["case_id"]
    old_path = doc_dict.get("file_path") or ""

    # Reuse path if it's the new-format relative path, otherwise allocate a fresh one
    if old_path and not old_path.startswith("tenants/"):
        file_path = Path(UPLOAD_BASE_DIR) / old_path
        file_path.parent.mkdir(parents=True, exist_ok=True)
        file_path.write_bytes(file_bytes)
        relative_path = old_path
    else:
        safe_filename = f"{doc_id}_{filename}"
        upload_dir = Path(UPLOAD_BASE_DIR) / tenant_id / case_id
        upload_dir.mkdir(parents=True, exist_ok=True)
        file_path = upload_dir / safe_filename
        file_path.write_bytes(file_bytes)
        relative_path = f"{tenant_id}/{case_id}/{safe_filename}"

    with get_db() as db:
        db.execute(
            "UPDATE documents SET filename = ?, file_path = ?, file_size = ?, mime_type = ? WHERE id = ?",
            (filename, relative_path, file_size, mime_type, doc_id)
        )
        updated = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
        updated_dict = dict(updated)

    # Background processing for evidence documents
    if doc_dict.get("category") == "evidence":
        if doc_dict.get("exhibit_label"):
            # Existing exhibit: preserve the label/name, just re-stamp with the new file
            background_tasks.add_task(
                _restamp_existing_exhibit, doc_id, case_id, tenant_id,
                file_bytes, filename, ext, file_path,
                doc_dict["exhibit_label"], doc_dict.get("exhibit_name") or ""
            )
        else:
            # No label yet: run full AI processing to assign one
            background_tasks.add_task(
                _process_exhibit_async, doc_id, case_id, tenant_id,
                file_bytes, filename, ext, file_path, relative_path, "evidence"
            )
            updated_dict["exhibit_processing"] = True

    return updated_dict


@router.patch("/documents/{doc_id}/category")
async def update_document_category(
    doc_id: str,
    req: DocumentCategoryUpdate,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Update a document's category. If changed TO 'evidence', triggers exhibit AI processing."""
    valid_categories = {'general', 'evidence', 'petition', 'correspondence', 'court_filing', 'ready'}
    cat = (req.category or "").strip().lower()
    if cat not in valid_categories:
        raise HTTPException(status_code=400, detail=f"Invalid category. Must be one of: {', '.join(sorted(valid_categories))}")

    with get_db() as db:
        doc = db.execute(
            "SELECT * FROM documents WHERE id = ? AND tenant_id = ?",
            (doc_id, current_user["tenant_id"])
        ).fetchone()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        old_category = doc["category"] or ""
        db.execute("UPDATE documents SET category = ? WHERE id = ?", (cat, doc_id))
        updated = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
        updated_dict = dict(updated)

    # If category changed TO evidence and the document has no exhibit_label yet, trigger AI processing
    if cat == "evidence" and old_category != "evidence" and not updated_dict.get("exhibit_label"):
        fp = Path(UPLOAD_BASE_DIR) / updated_dict["file_path"]
        if fp.exists():
            file_bytes = fp.read_bytes()
            ext = Path(updated_dict["file_path"]).suffix.lower()
            background_tasks.add_task(
                _process_exhibit_async, doc_id, updated_dict["case_id"], updated_dict["tenant_id"],
                file_bytes, updated_dict["filename"], ext, fp, updated_dict["file_path"],
                "evidence"
            )
            updated_dict["exhibit_processing"] = True

    return updated_dict


# --------------- Document merge endpoint ---------------

@router.post("/{case_id}/documents/merge")
async def merge_documents(
    case_id: str,
    req: DocumentMergeRequest,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Merge documents into a single PDF. Supports ALL file types.
    If merge_all=true, merges every document in the case. Otherwise merges doc_ids (1+)."""
    import fitz

    tenant_id = current_user["tenant_id"]

    with get_db() as db:
        # Verify case
        case = db.execute(
            "SELECT id FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, tenant_id)
        ).fetchone()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")

        if req.merge_all:
            # Merge ALL documents in the case
            rows = db.execute(
                "SELECT * FROM documents WHERE case_id = ? AND tenant_id = ? ORDER BY exhibit_order ASC, created_at ASC",
                (case_id, tenant_id)
            ).fetchall()
            if not rows:
                raise HTTPException(status_code=400, detail="No documents in this case to merge")
            docs = [dict(r) for r in rows]
            # Sort: court_filing FIRST (main document), then petition, then exhibits by label (A, B, C...)
            def _merge_sort_key(d):
                label = d.get('exhibit_label') or ''
                cat = d.get('category', '')
                if cat == 'court_filing':
                    return (0, '', d.get('created_at', ''))
                elif cat == 'petition':
                    return (1, '', d.get('created_at', ''))
                elif cat == 'evidence' and label:
                    return (2, label, d.get('created_at', ''))
                elif cat == 'evidence':
                    return (2, 'ZZZZ', d.get('created_at', ''))
                else:
                    return (3, '', d.get('created_at', ''))
            docs.sort(key=_merge_sort_key)
        else:
            if len(req.doc_ids) < 1:
                raise HTTPException(status_code=400, detail="Select at least 1 document to merge")
            # Fetch docs in the order requested
            docs = []
            for did in req.doc_ids:
                doc = db.execute(
                    "SELECT * FROM documents WHERE id = ? AND case_id = ? AND tenant_id = ?",
                    (did, case_id, tenant_id)
                ).fetchone()
                if not doc:
                    raise HTTPException(status_code=404, detail=f"Document {did} not found")
                docs.append(dict(doc))

    # Merge all file types using shared helper
    merged_pdf = fitz.open()
    for doc in docs:
        if not doc.get("file_path"):
            continue
        fp = Path(UPLOAD_BASE_DIR) / doc["file_path"]
        if not fp.exists():
            continue
        _add_file_to_pdf(merged_pdf, fp, doc.get("filename", ""))

    if merged_pdf.page_count == 0:
        merged_pdf.close()
        raise HTTPException(status_code=400, detail="No mergeable content found in selected documents")

    # Apply Bates numbering if requested
    if req.bates_prefix:
        _apply_bates_numbering(merged_pdf, req.bates_prefix, req.bates_start)

    # Save merged file
    merged_bytes = merged_pdf.tobytes()
    merged_pdf.close()

    doc_id = generate_id()
    upload_dir = Path(UPLOAD_BASE_DIR) / tenant_id / case_id
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_filename = f"{doc_id}_merged.pdf"
    file_path = upload_dir / safe_filename
    file_path.write_bytes(merged_bytes)
    relative_path = f"{tenant_id}/{case_id}/{safe_filename}"

    merged_count = len(docs)
    display_name = (req.filename or "Merged Document").strip()
    if not display_name.lower().endswith(".pdf"):
        display_name += ".pdf"

    with get_db() as db:
        db.execute(
            """INSERT INTO documents (id, case_id, tenant_id, filename, file_path, file_size, mime_type,
               category, uploaded_by, content_text, is_merged)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 1)""",
            (doc_id, case_id, tenant_id, display_name, relative_path, len(merged_bytes),
             "application/pdf", req.category, current_user["sub"],
             f"Merged from {merged_count} documents" + (f" | Bates: {req.bates_prefix}-{req.bates_start:06d}" if req.bates_prefix else ""))
        )

    # Skip AI renaming for merged documents — user already chose the name.
    # Only run exhibit processing if this is NOT a merge (individual uploads).
    # Merged documents keep their user-chosen display_name as-is.

    return {
        "id": doc_id,
        "case_id": case_id,
        "filename": display_name,
        "file_size": len(merged_bytes),
        "category": req.category,
        "merged_from": [d["id"] for d in docs],
        "merged_count": merged_count,
        "bates_applied": bool(req.bates_prefix),
        "exhibit_processing": req.category == "evidence",
    }


# --------------- Exhibit management endpoints ---------------

@router.patch("/documents/{doc_id}/exhibit")
async def update_exhibit(
    doc_id: str,
    req: ExhibitUpdate,
    current_user: dict = Depends(get_current_user)
):
    """Manually update exhibit name or label for a document."""
    with get_db() as db:
        doc = db.execute(
            "SELECT * FROM documents WHERE id = ? AND tenant_id = ?",
            (doc_id, current_user["tenant_id"])
        ).fetchone()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        updates = {}
        if req.exhibit_name is not None:
            updates["exhibit_name"] = req.exhibit_name
        if req.exhibit_label is not None:
            updates["exhibit_label"] = req.exhibit_label

        if not updates:
            return dict(doc)

        set_clause = ", ".join(f"{k} = ?" for k in updates.keys())
        values = list(updates.values()) + [doc_id]
        db.execute(f"UPDATE documents SET {set_clause} WHERE id = ?", values)

        # If label or name changed, update filename and re-stamp PDF
        if req.exhibit_label is not None or req.exhibit_name is not None:
            updated_doc = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
            updated_doc = dict(updated_doc)
            label = updated_doc.get("exhibit_label", "")
            name = updated_doc.get("exhibit_name", "")
            # Determine extension from file_path
            ext = Path(updated_doc.get("file_path", "")).suffix.lower() if updated_doc.get("file_path") else ".pdf"
            new_filename = f"Exhibit {label} — {name}{ext}"
            db.execute("UPDATE documents SET filename = ? WHERE id = ?", (new_filename, doc_id))

            if updated_doc.get("file_path") and updated_doc.get("mime_type", "").startswith("application/pdf"):
                try:
                    fp = Path(UPLOAD_BASE_DIR) / updated_doc["file_path"]
                    if fp.exists():
                        original_bytes = fp.read_bytes()
                        stamped = _stamp_exhibit_on_pdf(original_bytes, label, name)
                        fp.write_bytes(stamped)
                except Exception as e:
                    logger.warning(f"Re-stamp failed: {e}")

        updated = db.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
        return dict(updated)


@router.post("/{case_id}/documents/reorder")
async def reorder_exhibits(
    case_id: str,
    req: ExhibitReorderRequest,
    current_user: dict = Depends(get_current_user)
):
    """Reorder exhibits for a case. Reassigns labels based on the case numbering scheme."""
    tenant_id = current_user["tenant_id"]
    with get_db() as db:
        # Verify case and get numbering scheme
        case = db.execute(
            "SELECT id, exhibit_numbering FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, tenant_id)
        ).fetchone()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")

        numbering = (case["exhibit_numbering"] or "letters") if case else "letters"
        if numbering not in EXHIBIT_SCHEMES:
            numbering = "letters"

        # Sort items by requested order
        sorted_items = sorted(req.items, key=lambda x: x.order)

        for i, item in enumerate(sorted_items):
            new_label = _make_exhibit_label(i, numbering)
            doc = db.execute("SELECT * FROM documents WHERE id = ?", (item.doc_id,)).fetchone()
            exhibit_name = doc["exhibit_name"] if doc else ""
            ext = Path(doc["file_path"]).suffix.lower() if doc and doc.get("file_path") else ".pdf"
            new_filename = _exhibit_filename(new_label, exhibit_name, ext, numbering)

            db.execute(
                "UPDATE documents SET exhibit_order = ?, exhibit_label = ?, filename = ? WHERE id = ? AND case_id = ?",
                (i, new_label, new_filename, item.doc_id, case_id)
            )

            # Re-stamp PDFs with new label
            if doc and doc["file_path"] and (doc.get("mime_type") or "").startswith("application/pdf"):
                try:
                    fp = Path(UPLOAD_BASE_DIR) / doc["file_path"]
                    if fp.exists():
                        original_bytes = fp.read_bytes()
                        stamped = _stamp_exhibit_on_pdf(
                            original_bytes, new_label, exhibit_name
                        )
                        fp.write_bytes(stamped)
                except Exception as e:
                    logger.warning(f"Re-stamp on reorder failed for {item.doc_id}: {e}")

        # Return updated documents
        docs = db.execute(
            "SELECT * FROM documents WHERE case_id = ? ORDER BY exhibit_order ASC, created_at ASC",
            (case_id,)
        ).fetchall()
        return [dict(d) for d in docs]


@router.post("/documents/{doc_id}/retrigger-exhibit-ai")
async def retrigger_exhibit_ai(
    doc_id: str,
    background_tasks: BackgroundTasks,
    current_user: dict = Depends(get_current_user)
):
    """Re-trigger AI naming for a document."""
    with get_db() as db:
        doc = db.execute(
            "SELECT * FROM documents WHERE id = ? AND tenant_id = ?",
            (doc_id, current_user["tenant_id"])
        ).fetchone()
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        doc = dict(doc)
        if not doc.get("file_path"):
            raise HTTPException(status_code=400, detail="No file associated with this document")

        fp = Path(UPLOAD_BASE_DIR) / doc["file_path"]
        if not fp.exists():
            raise HTTPException(status_code=404, detail="File not found on server")

        file_bytes = fp.read_bytes()
        ext = Path(doc["filename"]).suffix.lower()

        background_tasks.add_task(
            _retrigger_ai_task, doc_id, doc["case_id"], doc["tenant_id"],
            file_bytes, doc["filename"], ext, fp
        )

    return {"status": "processing", "doc_id": doc_id}


def _retrigger_ai_task(doc_id: str, case_id: str, tenant_id: str,
                       file_bytes: bytes, filename: str, ext: str,
                       file_path_on_disk: Path):
    """Background: re-run AI naming and re-stamp PDF."""
    try:
        if ext == ".pdf":
            text = _extract_text_from_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            text = _extract_text_from_docx(file_bytes)
        elif ext == ".txt":
            text = file_bytes.decode("utf-8", errors="ignore")[:3000]
        else:
            text = ""

        exhibit_name = _ai_generate_exhibit_name(text, filename)

        with get_db() as db:
            doc = db.execute("SELECT exhibit_label, file_path FROM documents WHERE id = ?", (doc_id,)).fetchone()
            exhibit_label = doc["exhibit_label"] if doc else "A"
            file_ext = Path(doc["file_path"]).suffix.lower() if doc and doc.get("file_path") else ext
            new_filename = f"Exhibit {exhibit_label} — {exhibit_name}{file_ext}"

            db.execute(
                "UPDATE documents SET exhibit_name = ?, filename = ? WHERE id = ?",
                (exhibit_name, new_filename, doc_id)
            )

            if ext == ".pdf":
                stamped = _stamp_exhibit_on_pdf(file_bytes, exhibit_label, exhibit_name)
                file_path_on_disk.write_bytes(stamped)

        logger.info(f"AI re-trigger complete: {doc_id} -> {new_filename}")
    except Exception as e:
        logger.error(f"AI re-trigger failed for {doc_id}: {e}")


# USCIS status sub-route
@router.get("/{case_id}/uscis-status")
async def get_uscis_status(case_id: str, current_user: dict = Depends(get_current_user)):
    """Get USCIS case status for a case."""
    with get_db() as db:
        case = db.execute(
            "SELECT uscis_receipt_number, uscis_status, uscis_last_checked FROM cases WHERE id = ? AND tenant_id = ?",
            (case_id, current_user["tenant_id"])
        ).fetchone()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")
        return dict(case)
