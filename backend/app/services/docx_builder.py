"""
Professional DOCX Export Engine using python-docx.
Generates court-ready Word documents with proper formatting:
- Jurisdiction-specific captions (rendered as invisible-border tables for alignment)
- Proper page setup (US Letter / A4 per jurisdiction)
- Court-specific fonts, margins, line spacing
- Numbered paragraphs, signature blocks, exhibit tables
- Headers/footers with page numbers
"""
import re
import html as html_mod
from io import BytesIO
from typing import Optional
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── Color palette ────────────────────────────────────────────────
COLORS = {
    "primary": RGBColor(0x1F, 0x38, 0x64),      # Deep navy
    "secondary": RGBColor(0x2E, 0x40, 0x57),     # Dark slate
    "accent": RGBColor(0x2E, 0x75, 0xB6),        # Mid blue
    "text": RGBColor(0x00, 0x00, 0x00),           # Black
    "text_light": RGBColor(0x55, 0x55, 0x55),     # Grey
    "border": RGBColor(0xCC, 0xCC, 0xCC),         # Light grey
}


def build_docx(
    caption_data: dict,
    body_sections: list,
    court_rules: Optional[dict] = None,
    financials: Optional[dict] = None,
) -> bytes:
    """
    Build a complete court-ready DOCX document.

    Args:
        caption_data: Output from caption_engine.generate_caption_for_docx()
        body_sections: List of section dicts from Claude's structured JSON
        court_rules: Optional court formatting rules
        financials: Optional financial calculation results

    Returns:
        DOCX file as bytes
    """
    doc = Document()

    # ─── Page setup ───────────────────────────────────────────
    _setup_page(doc, court_rules, caption_data.get("style_key", "us_federal"))

    # ─── Default font ─────────────────────────────────────────
    font_name = "Times New Roman"
    font_size = 12
    line_spacing = 2.0
    if court_rules:
        font_name = court_rules.get("default_font", "Times New Roman")
        font_size = court_rules.get("font_size", 12)
        line_spacing = court_rules.get("line_spacing", 2.0)

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.font.color.rgb = COLORS["text"]
    style.paragraph_format.line_spacing = line_spacing
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ─── Header and Footer ───────────────────────────────────
    _add_header_footer(doc, caption_data)

    # ─── Caption ──────────────────────────────────────────────
    _render_caption(doc, caption_data, font_name, font_size)

    # ─── Body sections ────────────────────────────────────────
    _render_body(doc, body_sections, font_name, font_size, line_spacing, financials)

    # ─── Save to bytes ────────────────────────────────────────
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _setup_page(doc: Document, court_rules: Optional[dict], style_key: str):
    """Configure page size and margins."""
    section = doc.sections[0]

    # US Letter by default (8.5 x 11 inches)
    # A4 for UK, India, Australia, Hong Kong, Ireland
    a4_jurisdictions = {"uk", "india", "australia", "hong_kong", "ireland"}
    if style_key in a4_jurisdictions:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # Margins
    if court_rules:
        section.top_margin = Inches(court_rules.get("margin_top", 1.0))
        section.bottom_margin = Inches(court_rules.get("margin_bottom", 1.0))
        section.left_margin = Inches(court_rules.get("margin_left", 1.0))
        section.right_margin = Inches(court_rules.get("margin_right", 1.0))
    else:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)


def _add_header_footer(doc: Document, caption_data: dict):
    """Add header with case title and footer with page numbers."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Regular header (not on first page)
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(caption_data.get("document_title", ""))
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["text_light"]
    run.font.italic = True

    # Add bottom border to header
    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="1F3864"/></w:pBdr>')
    pPr.append(pBdr)

    # Footer with page numbers
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add top border to footer
    pPr = fp._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="4" w:color="1F3864"/></w:pBdr>')
    pPr.append(pBdr)

    run = fp.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["text_light"]

    # Current page number field
    fld_xml = f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
    fp._p.append(parse_xml(fld_xml))

    run = fp.add_run(" of ")
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["text_light"]

    # Total pages field
    fld_xml2 = f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
    fp._p.append(parse_xml(fld_xml2))

    # First page footer (same page numbers)
    first_footer = section.first_page_footer
    ffp = first_footer.paragraphs[0] if first_footer.paragraphs else first_footer.add_paragraph()
    ffp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ffp.add_run("Page ")
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["text_light"]
    ffp._p.append(parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'))
    run = ffp.add_run(" of ")
    run.font.size = Pt(9)
    run.font.color.rgb = COLORS["text_light"]
    ffp._p.append(parse_xml(f'<w:fldSimple {nsdecls("w")} w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'))


def _render_caption(doc: Document, caption_data: dict, font_name: str, font_size: int):
    """Render the caption/header section using invisible-border tables for alignment."""
    party_format = caption_data.get("party_format", "parentheses")

    # Court heading — centered, bold, uppercase
    for line in caption_data.get("court_lines", []):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line.upper())
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = font_name
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

    # Case number (for non-US styles, shown below court heading)
    if party_format != "parentheses" and caption_data.get("case_number_text"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption_data["case_number_text"])
        run.bold = True
        run.font.size = Pt(font_size)
        run.font.name = font_name
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    # "In the matter of" line
    if caption_data.get("in_the_matter_of"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"IN THE MATTER OF {caption_data['in_the_matter_of'].upper()}")
        run.bold = True
        run.font.size = Pt(font_size - 1)
        run.font.name = font_name

    # Divider
    _add_divider(doc)

    # Party block
    if party_format == "parentheses":
        _render_us_parties(doc, caption_data, font_name, font_size)
    elif party_format == "between_and_dotted":
        _render_commonwealth_parties(doc, caption_data, font_name, font_size, dotted=True)
    else:
        _render_commonwealth_parties(doc, caption_data, font_name, font_size, dotted=False)

    # Divider
    _add_divider(doc)

    # Document title
    if caption_data.get("document_title"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption_data["document_title"])
        run.bold = True
        run.font.size = Pt(font_size + 2)
        run.font.name = font_name
        run.underline = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(18)

    # Extra spacing before body
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _render_us_parties(doc: Document, caption_data: dict, font_name: str, font_size: int):
    """Render US-style parties with parentheses and case number on right."""
    left = caption_data.get("left_parties", [])
    right = caption_data.get("right_parties", [])
    case_num = caption_data.get("case_number_text", "")

    # Use a 3-column table: party name | ) | case number
    col_count = 3
    row_count = len(left) * 2 + 3 + len(right) * 2  # name+role per party, v. row, spacers
    table = doc.add_table(rows=0, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove table borders
    _set_table_borders(table, none=True)

    # Set column widths
    for cell_idx in range(col_count):
        for row in table.rows:
            pass  # Will set on cells as we add rows

    # Left parties
    for i, p in enumerate(left):
        row = table.add_row()
        _set_cell_text(row.cells[0], f"{p['name']},", font_name, font_size, bold=True)
        _set_cell_text(row.cells[1], ")", font_name, font_size, align="center")
        _set_cell_text(row.cells[2], "", font_name, font_size)

        row2 = table.add_row()
        _set_cell_text(row2.cells[0], f"     {p['role']},", font_name, font_size, italic=True)
        _set_cell_text(row2.cells[1], ")", font_name, font_size, align="center")
        _set_cell_text(row2.cells[2], "", font_name, font_size)

    # Blank + v. + case number
    row_blank = table.add_row()
    _set_cell_text(row_blank.cells[0], "", font_name, font_size)
    _set_cell_text(row_blank.cells[1], ")", font_name, font_size, align="center")
    _set_cell_text(row_blank.cells[2], case_num, font_name, font_size, bold=True, align="right")

    row_v = table.add_row()
    _set_cell_text(row_v.cells[0], "v.", font_name, font_size, bold=True)
    _set_cell_text(row_v.cells[1], ")", font_name, font_size, align="center")
    _set_cell_text(row_v.cells[2], "", font_name, font_size)

    row_blank2 = table.add_row()
    _set_cell_text(row_blank2.cells[0], "", font_name, font_size)
    _set_cell_text(row_blank2.cells[1], ")", font_name, font_size, align="center")
    _set_cell_text(row_blank2.cells[2], "", font_name, font_size)

    # Right parties
    for i, p in enumerate(right):
        row = table.add_row()
        _set_cell_text(row.cells[0], f"{p['name']},", font_name, font_size, bold=True)
        _set_cell_text(row.cells[1], ")", font_name, font_size, align="center")
        _set_cell_text(row.cells[2], "", font_name, font_size)

        period = "." if i == len(right) - 1 else ","
        row2 = table.add_row()
        _set_cell_text(row2.cells[0], f"     {p['role']}{period}", font_name, font_size, italic=True)
        _set_cell_text(row2.cells[1], ")", font_name, font_size, align="center")
        _set_cell_text(row2.cells[2], "", font_name, font_size)

    # Set column widths after all rows are added
    for row in table.rows:
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(0.5)
        row.cells[2].width = Inches(2.5)


def _render_commonwealth_parties(doc: Document, caption_data: dict, font_name: str, font_size: int, dotted: bool = True):
    """Render Commonwealth-style BETWEEN/AND parties."""
    left = caption_data.get("left_parties", [])
    right = caption_data.get("right_parties", [])

    # "BETWEEN" header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BETWEEN")
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = font_name
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    # Left parties
    for party in left:
        if dotted:
            # Use a 2-column table for dotted leader alignment
            t = doc.add_table(rows=1, cols=2)
            _set_table_borders(t, none=True)
            _set_cell_text(t.rows[0].cells[0], party["name"], font_name, font_size, bold=True)
            _set_cell_text(t.rows[0].cells[1], party["role"].upper(), font_name, font_size, bold=True, align="right")
            t.rows[0].cells[0].width = Inches(4.0)
            t.rows[0].cells[1].width = Inches(2.5)
            # Add dotted border on bottom of first cell
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(party["name"])
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = font_name
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(party["role"])
            run2.italic = True
            run2.font.size = Pt(font_size)
            run2.font.name = font_name

    # AND separator
    sep = "AND" if dotted else "-and-"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(sep)
    run.bold = True
    run.font.size = Pt(font_size)
    run.font.name = font_name
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)

    # Right parties
    for party in right:
        if dotted:
            t = doc.add_table(rows=1, cols=2)
            _set_table_borders(t, none=True)
            _set_cell_text(t.rows[0].cells[0], party["name"], font_name, font_size, bold=True)
            _set_cell_text(t.rows[0].cells[1], party["role"].upper(), font_name, font_size, bold=True, align="right")
            t.rows[0].cells[0].width = Inches(4.0)
            t.rows[0].cells[1].width = Inches(2.5)
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(party["name"])
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = font_name
            p2 = doc.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = p2.add_run(party["role"])
            run2.italic = True
            run2.font.size = Pt(font_size)
            run2.font.name = font_name


def _render_body(doc: Document, sections: list, font_name: str, font_size: int, line_spacing: float, financials: Optional[dict]):
    """Render document body sections from Claude's structured JSON."""
    from app.services.interest_calc import substitute_financial_tokens

    numbered_counter = 0

    for section in sections:
        sec_type = section.get("type", "paragraph")
        content = section.get("content", "")

        # Substitute financial tokens if present
        if financials and content:
            content = substitute_financial_tokens(content, financials)

        if sec_type == "heading1":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(content.upper())
            run.bold = True
            run.font.size = Pt(font_size + 1)
            run.font.name = font_name
            run.font.color.rgb = COLORS["primary"]
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            # Add bottom border
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="1F3864"/></w:pBdr>')
            pPr.append(pBdr)
            numbered_counter = 0  # Reset numbering per section

        elif sec_type == "heading2":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(font_size)
            run.font.name = font_name
            run.font.color.rgb = COLORS["secondary"]
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)

        elif sec_type == "heading3":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.underline = True
            run.font.size = Pt(font_size)
            run.font.name = font_name
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)

        elif sec_type == "numberedParagraph":
            numbered_counter += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

            # Add number
            runs = section.get("runs")
            if runs:
                num_run = p.add_run(f"{numbered_counter}. ")
                num_run.font.size = Pt(font_size)
                num_run.font.name = font_name
                for r in runs:
                    text = r.get("text", "")
                    if financials:
                        text = substitute_financial_tokens(text, financials)
                    run = p.add_run(text)
                    run.font.size = Pt(font_size)
                    run.font.name = font_name
                    run.bold = r.get("bold", False)
                    run.italic = r.get("italic", False)
            else:
                run = p.add_run(f"{numbered_counter}. {content}")
                run.font.size = Pt(font_size)
                run.font.name = font_name

        elif sec_type == "paragraph":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_after = Pt(6)

            runs = section.get("runs")
            if runs:
                for r in runs:
                    text = r.get("text", "")
                    if financials:
                        text = substitute_financial_tokens(text, financials)
                    run = p.add_run(text)
                    run.font.size = Pt(font_size)
                    run.font.name = font_name
                    run.bold = r.get("bold", False)
                    run.italic = r.get("italic", False)
                    if r.get("exhibitRef"):
                        run.bold = True
                        run.italic = True
            else:
                run = p.add_run(content)
                run.font.size = Pt(font_size)
                run.font.name = font_name

        elif sec_type == "bullet":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.75)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            run = p.add_run(f"\u2022  {content}")
            run.font.size = Pt(font_size)
            run.font.name = font_name
            p.paragraph_format.space_after = Pt(3)

        elif sec_type == "table":
            table_data = section.get("tableData", {})
            if table_data:
                _render_table(doc, table_data, font_name, financials)

        elif sec_type == "signatureBlock":
            _render_signature_block(doc, section, font_name, font_size)

        elif sec_type == "exhibitTable":
            exhibits = section.get("exhibits", [])
            if exhibits:
                _render_exhibit_table(doc, exhibits, font_name, font_size)

        elif sec_type == "spacer":
            doc.add_paragraph().paragraph_format.space_after = Pt(section.get("points", 12))

        elif sec_type == "divider":
            _add_divider(doc)

        elif sec_type == "pageBreak":
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_break(docx_enum=7)  # WD_BREAK.PAGE


def _render_table(doc: Document, table_data: dict, font_name: str, financials: Optional[dict]):
    """Render a data table."""
    from app.services.interest_calc import substitute_financial_tokens

    rows_data = table_data.get("rows", [])
    headers = table_data.get("headers", [])
    if not rows_data and not headers:
        return

    # Normalize rows — html_to_sections produces plain lists, structured input uses dicts with "cells"
    normalized_rows = []
    if headers:
        normalized_rows.append({"cells": headers, "isHeader": True})
    for rd in rows_data:
        if isinstance(rd, list):
            normalized_rows.append({"cells": rd, "isHeader": False})
        elif isinstance(rd, dict):
            normalized_rows.append(rd)

    if not normalized_rows:
        return

    col_count = len(normalized_rows[0].get("cells", []))
    if col_count == 0:
        return
    table = doc.add_table(rows=0, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_data in normalized_rows:
        row = table.add_row()
        is_header = row_data.get("isHeader", False)
        shading = row_data.get("shading")

        for i, cell_text in enumerate(row_data.get("cells", [])):
            cell = row.cells[i]
            if financials:
                cell_text = substitute_financial_tokens(cell_text, financials)

            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(10)
            run.font.name = font_name
            run.bold = is_header

            # Cell margins
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:left w:w="120" w:type="dxa"/><w:right w:w="120" w:type="dxa"/></w:tcMar>')
            tcPr.append(tcMar)

            # Header shading
            if is_header:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864" w:val="clear"/>')
                tcPr.append(shd)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            elif shading:
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{shading}" w:val="clear"/>')
                tcPr.append(shd)

    # Set table borders
    _set_table_borders(table, none=False)


def _render_signature_block(doc: Document, section: dict, font_name: str, font_size: int):
    """Render a signature block."""
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    content = section.get("content", "")
    lines = content.split("\n") if content else [
        "Respectfully submitted,",
        "",
        "____________________________",
        "[Attorney Name]",
        "[Bar Number]",
        "[Firm Name]",
        "[Address]",
        "[Phone]",
        "[Email]",
        f"Date: {datetime.now().strftime('%B %d, %Y')}",
    ]

    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Inches(3.5)
        run = p.add_run(line)
        run.font.size = Pt(font_size)
        run.font.name = font_name
        p.paragraph_format.space_after = Pt(0)


def _render_exhibit_table(doc: Document, exhibits: list, font_name: str, font_size: int):
    """Render an exhibit list table."""
    p = doc.add_paragraph()
    run = p.add_run("EXHIBIT LIST")
    run.bold = True
    run.font.size = Pt(font_size + 1)
    run.font.name = font_name
    run.font.color.rgb = COLORS["primary"]
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0]
    _set_cell_text(hdr.cells[0], "Exhibit", font_name, 10, bold=True)
    _set_cell_text(hdr.cells[1], "Description", font_name, 10, bold=True)

    # Apply header shading
    for cell in hdr.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F3864" w:val="clear"/>')
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Rows
    for i, exhibit in enumerate(exhibits):
        row = table.add_row()
        _set_cell_text(row.cells[0], exhibit.get("label", f"Exhibit {chr(65 + i)}"), font_name, 10, bold=True)
        _set_cell_text(row.cells[1], exhibit.get("description", ""), font_name, 10)

        # Alternating row shading
        if i % 2 == 0:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F9" w:val="clear"/>')
                tcPr.append(shd)

    hdr.cells[0].width = Inches(1.5)
    hdr.cells[1].width = Inches(5.0)
    _set_table_borders(table, none=False)


# ─── Helper functions ─────────────────────────────────────────────

def _add_divider(doc: Document):
    """Add a horizontal divider line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def _set_cell_text(cell, text: str, font_name: str, font_size: int, bold: bool = False, italic: bool = False, align: str = "left"):
    """Set text in a table cell with formatting."""
    p = cell.paragraphs[0]
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def _set_table_borders(table, none: bool = False):
    """Set or remove table borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    if none:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tblBorders>'
        )
    else:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
            '</w:tblBorders>'
        )

    # Remove existing borders element if present
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


# ─── HTML-to-Sections Parser ─────────────────────────────────────
def html_to_sections(html_content: str) -> list:
    """
    Parse HTML content (from Claude AI output) into the structured section format
    that _render_body() expects. Handles: h2, h3, p, blockquote, ol/ul/li, table, strong, em.
    """
    from html.parser import HTMLParser

    sections = []
    current_runs = []
    current_tag_stack = []
    current_text = []
    in_table = False
    table_headers = []
    table_rows = []
    current_row = []
    current_cell_text = []
    in_list = False
    list_type = "ul"
    list_items = []
    current_li_text = []

    class LegalHTMLParser(HTMLParser):
        def handle_starttag(self, tag, attrs):
            nonlocal in_table, table_headers, table_rows, current_row, current_cell_text
            nonlocal in_list, list_type, list_items, current_li_text

            tag = tag.lower()

            if tag in ("h2", "h3"):
                _flush_paragraph()
                current_tag_stack.append(tag)
            elif tag == "p":
                _flush_paragraph()
                current_tag_stack.append("p")
            elif tag == "blockquote":
                _flush_paragraph()
                current_tag_stack.append("blockquote")
            elif tag in ("strong", "b"):
                _flush_text()
                current_tag_stack.append("bold")
            elif tag in ("em", "i"):
                _flush_text()
                current_tag_stack.append("italic")
            elif tag in ("ol", "ul"):
                _flush_paragraph()
                in_list = True
                list_type = tag
                list_items = []
            elif tag == "li":
                current_li_text = []
            elif tag == "table":
                _flush_paragraph()
                in_table = True
                table_headers = []
                table_rows = []
            elif tag == "tr":
                current_row = []
            elif tag in ("th", "td"):
                current_cell_text = []
            elif tag == "br":
                current_text.append("\n")

        def handle_endtag(self, tag):
            nonlocal in_table, table_headers, table_rows, current_row, current_cell_text
            nonlocal in_list, list_type, list_items, current_li_text

            tag = tag.lower()

            if tag in ("h2", "h3"):
                _flush_text()
                text = "".join(t.get("text", "") for t in current_runs).strip()
                if tag == "h2":
                    sections.append({"type": "heading1", "content": text})
                else:
                    sections.append({"type": "heading2", "content": text})
                current_runs.clear()
                if current_tag_stack and current_tag_stack[-1] == tag:
                    current_tag_stack.pop()
            elif tag == "p":
                _flush_text()
                text = "".join(t.get("text", "") for t in current_runs).strip()
                if text:
                    # Check if it starts with a number (numbered paragraph)
                    stripped = text.lstrip()
                    if stripped and stripped[0].isdigit() and ". " in stripped[:6]:
                        # Remove the leading number — _render_body adds its own
                        dot_idx = stripped.index(". ")
                        body_text = stripped[dot_idx + 2:]
                        sections.append({
                            "type": "numberedParagraph",
                            "content": body_text,
                            "runs": [{"text": body_text}] if len(current_runs) <= 1 else current_runs[:],
                        })
                    else:
                        sections.append({
                            "type": "paragraph",
                            "content": text,
                            "runs": current_runs[:] if len(current_runs) > 1 else None,
                        })
                current_runs.clear()
                if current_tag_stack and current_tag_stack[-1] == "p":
                    current_tag_stack.pop()
            elif tag == "blockquote":
                _flush_text()
                text = "".join(t.get("text", "") for t in current_runs).strip()
                if text:
                    sections.append({"type": "paragraph", "content": text, "runs": [
                        {"text": text, "italic": True}
                    ]})
                current_runs.clear()
                if current_tag_stack and current_tag_stack[-1] == "blockquote":
                    current_tag_stack.pop()
            elif tag in ("strong", "b"):
                _flush_text()
                if current_tag_stack and current_tag_stack[-1] == "bold":
                    current_tag_stack.pop()
            elif tag in ("em", "i"):
                _flush_text()
                if current_tag_stack and current_tag_stack[-1] == "italic":
                    current_tag_stack.pop()
            elif tag == "li":
                li_text = "".join(current_li_text).strip()
                if li_text:
                    list_items.append(li_text)
            elif tag in ("ol", "ul"):
                for item in list_items:
                    sections.append({"type": "bullet", "content": item})
                in_list = False
                list_items = []
            elif tag in ("th",):
                cell_text = "".join(current_cell_text).strip()
                table_headers.append(cell_text)
            elif tag == "td":
                cell_text = "".join(current_cell_text).strip()
                current_row.append(cell_text)
            elif tag == "tr":
                if current_row:
                    table_rows.append(current_row)
                current_row = []
            elif tag == "table":
                if table_headers or table_rows:
                    sections.append({
                        "type": "table",
                        "tableData": {
                            "headers": table_headers if table_headers else (table_rows[0] if table_rows else []),
                            "rows": table_rows if table_headers else table_rows[1:],
                        }
                    })
                in_table = False
                table_headers = []
                table_rows = []

        def handle_data(self, data):
            nonlocal current_cell_text, current_li_text
            if in_table:
                current_cell_text.append(data)
                return
            if in_list:
                current_li_text.append(data)
                return
            current_text.append(data)

        def handle_entityref(self, name):
            char = html_mod.unescape(f"&{name};")
            self.handle_data(char)

        def handle_charref(self, name):
            char = html_mod.unescape(f"&#{name};")
            self.handle_data(char)

    def _flush_text():
        text = "".join(current_text).strip()
        if text:
            is_bold = "bold" in current_tag_stack
            is_italic = "italic" in current_tag_stack
            current_runs.append({"text": text, "bold": is_bold, "italic": is_italic})
        current_text.clear()

    def _flush_paragraph():
        _flush_text()
        text = "".join(t.get("text", "") for t in current_runs).strip()
        if text:
            sections.append({"type": "paragraph", "content": text, "runs": current_runs[:]})
        current_runs.clear()

    parser = LegalHTMLParser()
    parser.feed(html_content)
    _flush_paragraph()  # flush any remaining content

    return sections
